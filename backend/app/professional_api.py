"""Database-backed adult research, assessment, career and institution APIs."""

from __future__ import annotations

import json
import re
import textwrap
import uuid
import zipfile
from datetime import datetime, timezone
from io import BytesIO
from urllib.parse import quote
from urllib.request import Request, urlopen

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from sqlalchemy import func, select

from app import ai_tutor, local_library
from app.database import session_scope
from app.models import (
    Assessment,
    Assignment,
    AttendanceRecord,
    Attempt,
    AuditEvent,
    Cohort,
    Course,
    CPDRecord,
    DiscussionPost,
    LearningItem,
    Membership,
    Organization,
    PortfolioItem,
    Question,
    ResearchDocument,
    ResearchProject,
    Resource,
    Submission,
    User,
)

router = APIRouter(prefix="/api/pro", tags=["professional"])


def _as_dict(row, *fields):
    return {field: getattr(row, field) for field in fields}


def _required(body: dict, key: str) -> str:
    value = str(body.get(key, "")).strip()
    if not value:
        raise HTTPException(400, f"{key} is required")
    return value


def _audit(session, actor_id, action, entity_type, entity_id, after=None, before=None):
    session.add(AuditEvent(
        actor_id=actor_id or None,
        action=action,
        entity_type=entity_type,
        entity_id=entity_id,
        before=before,
        after=after,
    ))


@router.post("/users/ensure")
def ensure_user(body: dict):
    email = _required(body, "email").lower()
    if email != "shovan@legacy.local":
        raise HTTPException(403, "EduAI_Pro is configured for the single administrator Shovan")
    with session_scope() as session:
        user = session.scalar(select(User).where(User.email == email))
        if not user:
            user = User(
                email=email,
                display_name="Shovan",
                role="admin",
            )
            session.add(user)
            session.flush()
            _audit(session, user.id, "create", "user", user.id, {"email": email})
        return _as_dict(user, "id", "email", "display_name", "role", "settings")


@router.get("/dashboard/{user_id}")
def dashboard(user_id: str):
    with session_scope() as session:
        if not session.get(User, user_id):
            raise HTTPException(404, "User not found")
        counts = {
            "research_projects": session.scalar(select(func.count()).select_from(ResearchProject).where(ResearchProject.owner_id == user_id)),
            "notes": session.scalar(select(func.count()).select_from(LearningItem).where(LearningItem.user_id == user_id)),
            "portfolio_items": session.scalar(select(func.count()).select_from(PortfolioItem).where(PortfolioItem.owner_id == user_id)),
            "cpd_hours": session.scalar(select(func.coalesce(func.sum(CPDRecord.hours), 0)).where(CPDRecord.owner_id == user_id)),
            "assessment_attempts": session.scalar(select(func.count()).select_from(Attempt).where(Attempt.user_id == user_id)),
        }
        return counts


# Research and knowledge -------------------------------------------------------

@router.get("/research/projects/{user_id}")
def list_projects(user_id: str):
    with session_scope() as session:
        rows = session.scalars(
            select(ResearchProject).where(ResearchProject.owner_id == user_id, ResearchProject.deleted_at.is_(None))
        ).all()
        return [_as_dict(row, "id", "title", "description", "settings", "created_at") for row in rows]


@router.post("/research/projects")
def create_project(body: dict):
    with session_scope() as session:
        project = ResearchProject(
            owner_id=_required(body, "owner_id"),
            title=_required(body, "title"),
            description=str(body.get("description", "")),
            settings=body.get("settings") or {},
        )
        session.add(project)
        session.flush()
        _audit(session, project.owner_id, "create", "research_project", project.id, {"title": project.title})
        return _as_dict(project, "id", "title", "description", "settings")


@router.post("/research/projects/{project_id}/local/{file_id}")
def import_local_document(project_id: str, file_id: str):
    result = local_library.get_file(file_id)
    if not result:
        raise HTTPException(404, "Indexed local file not found")
    record, path = result
    text = local_library.extract_text(path)
    with session_scope() as session:
        project = session.get(ResearchProject, project_id)
        if not project:
            raise HTTPException(404, "Research project not found")
        resource = Resource(
            owner_id=project.owner_id,
            kind=record["category"],
            title=record["filename"],
            local_path=str(path),
            extracted_text=text,
            citation={"title": record["filename"], "type": "owned_local_file"},
            metadata_json={"local_file_id": file_id},
        )
        session.add(resource)
        session.flush()
        document = ResearchDocument(
            project_id=project_id,
            resource_id=resource.id,
            title=resource.title,
            citation=resource.citation,
            analysis=local_library.analyse_text(text),
        )
        session.add(document)
        session.flush()
        _audit(session, project.owner_id, "import", "research_document", document.id, {"path": str(path)})
        return _as_dict(document, "id", "project_id", "resource_id", "title", "citation", "analysis")


@router.get("/research/projects/{project_id}/documents")
def list_project_documents(project_id: str):
    with session_scope() as session:
        rows = session.scalars(select(ResearchDocument).where(ResearchDocument.project_id == project_id)).all()
        return [_as_dict(row, "id", "resource_id", "title", "citation", "analysis") for row in rows]


def _words(value: str):
    return {word.lower() for word in re.findall(r"[A-Za-z][A-Za-z0-9-]{2,}", value)}


@router.get("/research/search")
def research_search(q: str, project_id: str = "", limit: int = 20):
    wanted = _words(q)
    with session_scope() as session:
        statement = select(Resource).where(Resource.deleted_at.is_(None))
        if project_id:
            resource_ids = select(ResearchDocument.resource_id).where(ResearchDocument.project_id == project_id)
            statement = statement.where(Resource.id.in_(resource_ids))
        resources = session.scalars(statement.limit(3000)).all()
        scored = []
        for resource in resources:
            score = len(wanted & _words(f"{resource.title} {resource.extracted_text[:12000]}"))
            if score:
                scored.append((score, resource))
        scored.sort(key=lambda item: item[0], reverse=True)
        return [
            {
                "id": resource.id,
                "title": resource.title,
                "excerpt": resource.extracted_text[:500],
                "citation": resource.citation,
                "score": score,
            }
            for score, resource in scored[:max(1, min(limit, 100))]
        ]


@router.post("/research/notes")
def save_research_note(body: dict):
    with session_scope() as session:
        item = LearningItem(
            user_id=_required(body, "user_id"),
            kind="research_note",
            title=_required(body, "title"),
            content={
                "project_id": body.get("project_id"),
                "markdown": str(body.get("markdown", "")),
                "citations": body.get("citations") or [],
                "tags": body.get("tags") or [],
            },
        )
        session.add(item)
        session.flush()
        _audit(session, item.user_id, "create", "research_note", item.id, {"title": item.title})
        return _as_dict(item, "id", "title", "content", "created_at")


@router.get("/research/notes/{user_id}")
def list_research_notes(user_id: str):
    with session_scope() as session:
        rows = session.scalars(
            select(LearningItem).where(LearningItem.user_id == user_id, LearningItem.kind == "research_note")
        ).all()
        return [_as_dict(row, "id", "title", "content", "created_at", "updated_at") for row in rows]


@router.get("/research/doi/{doi:path}")
def doi_lookup(doi: str):
    try:
        request = Request(
            f"https://api.crossref.org/works/{quote(doi, safe='')}",
            headers={"User-Agent": "EduAI-Pro/1.0 (mailto:admin@localhost)"},
        )
        with urlopen(request, timeout=10) as response:
            message = json.loads(response.read().decode("utf-8"))["message"]
        return {
            "doi": message.get("DOI"),
            "title": (message.get("title") or [""])[0],
            "authors": [
                " ".join(filter(None, (author.get("given"), author.get("family"))))
                for author in message.get("author", [])
            ],
            "publisher": message.get("publisher"),
            "url": message.get("URL"),
            "issued": message.get("issued"),
        }
    except Exception as exc:
        raise HTTPException(502, f"DOI lookup failed: {type(exc).__name__}") from exc


@router.get("/research/projects/{project_id}/bibliography")
def bibliography(project_id: str, format: str = "markdown"):
    with session_scope() as session:
        docs = session.scalars(select(ResearchDocument).where(ResearchDocument.project_id == project_id)).all()
    if format == "bibtex":
        entries = []
        for index, doc in enumerate(docs, start=1):
            citation = doc.citation or {}
            entries.append(
                "@misc{eduai" + str(index) + ",\n"
                f"  title = {{{citation.get('title') or doc.title}}},\n"
                f"  author = {{{citation.get('author') or 'Unknown'}}},\n"
                f"  url = {{{citation.get('url') or ''}}}\n}}"
            )
        content, media, suffix = "\n\n".join(entries), "application/x-bibtex", "bib"
    else:
        content = "\n".join(
            f"- {(doc.citation or {}).get('author', 'Unknown')}. *{(doc.citation or {}).get('title', doc.title)}*. "
            f"{(doc.citation or {}).get('url', '')}"
            for doc in docs
        )
        media, suffix = "text/markdown", "md"
    return StreamingResponse(
        BytesIO(content.encode("utf-8")),
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="bibliography.{suffix}"'},
    )


# Assessments and assignments -------------------------------------------------

@router.get("/assessments")
def list_assessments(course_id: str = ""):
    with session_scope() as session:
        statement = select(Assessment).where(Assessment.deleted_at.is_(None))
        if course_id:
            statement = statement.where(Assessment.course_id == course_id)
        rows = session.scalars(statement).all()
        return [_as_dict(row, "id", "course_id", "title", "kind", "settings", "rubric") for row in rows]


@router.post("/assessments")
def create_assessment(body: dict):
    with session_scope() as session:
        assessment = Assessment(
            course_id=body.get("course_id") or None,
            owner_id=body.get("owner_id") or None,
            title=_required(body, "title"),
            kind=str(body.get("kind") or "quiz"),
            settings=body.get("settings") or {},
            rubric=body.get("rubric") or {},
        )
        session.add(assessment)
        session.flush()
        for question_data in body.get("questions") or []:
            session.add(Question(
                assessment_id=assessment.id,
                kind=str(question_data.get("kind") or "multiple_choice"),
                prompt=str(question_data.get("prompt") or ""),
                options=question_data.get("options") or [],
                answer=question_data.get("answer") or {},
                points=float(question_data.get("points") or 1),
            ))
        _audit(session, assessment.owner_id, "create", "assessment", assessment.id, {"title": assessment.title})
        return _as_dict(assessment, "id", "title", "kind", "settings", "rubric")


@router.post("/assessments/{assessment_id}/attempts")
def submit_attempt(assessment_id: str, body: dict):
    with session_scope() as session:
        questions = session.scalars(select(Question).where(Question.assessment_id == assessment_id)).all()
        answers = body.get("answers") or {}
        earned = 0.0
        possible = sum(question.points for question in questions)
        details = []
        for question in questions:
            expected = question.answer or {}
            given = answers.get(question.id)
            correct = given == expected.get("value")
            if correct:
                earned += question.points
            details.append({"question_id": question.id, "correct": correct})
        score = (earned / possible * 100) if possible else None
        attempt = Attempt(
            assessment_id=assessment_id,
            user_id=_required(body, "user_id"),
            answers=answers,
            score=score,
            feedback={"items": details},
            submitted_at=datetime.now(timezone.utc),
        )
        session.add(attempt)
        session.flush()
        return _as_dict(attempt, "id", "assessment_id", "user_id", "score", "feedback", "submitted_at")


@router.get("/assignments")
def list_assignments(course_id: str = ""):
    with session_scope() as session:
        statement = select(Assignment).where(Assignment.deleted_at.is_(None))
        if course_id:
            statement = statement.where(Assignment.course_id == course_id)
        rows = session.scalars(statement).all()
        return [_as_dict(row, "id", "course_id", "title", "instructions", "rubric", "due_at") for row in rows]


@router.post("/assignments")
def create_assignment(body: dict):
    with session_scope() as session:
        assignment = Assignment(
            course_id=_required(body, "course_id"),
            owner_id=_required(body, "owner_id"),
            title=_required(body, "title"),
            instructions=str(body.get("instructions", "")),
            rubric=body.get("rubric") or {},
            due_at=datetime.fromisoformat(body["due_at"]) if body.get("due_at") else None,
        )
        session.add(assignment)
        session.flush()
        _audit(session, assignment.owner_id, "create", "assignment", assignment.id)
        return _as_dict(assignment, "id", "course_id", "title", "instructions", "rubric", "due_at")


@router.post("/assignments/{assignment_id}/submissions")
def submit_assignment(assignment_id: str, body: dict):
    with session_scope() as session:
        submission = Submission(
            assignment_id=assignment_id,
            user_id=_required(body, "user_id"),
            body=str(body.get("body", "")),
            attachments=body.get("attachments") or [],
        )
        session.add(submission)
        session.flush()
        return _as_dict(submission, "id", "assignment_id", "user_id", "body", "attachments", "grade", "feedback")


# Career and professional records --------------------------------------------

@router.get("/portfolio/{user_id}")
def list_portfolio(user_id: str):
    with session_scope() as session:
        rows = session.scalars(select(PortfolioItem).where(PortfolioItem.owner_id == user_id)).all()
        return [_as_dict(row, "id", "title", "description", "evidence", "skills", "published") for row in rows]


@router.post("/portfolio")
def create_portfolio_item(body: dict):
    with session_scope() as session:
        item = PortfolioItem(
            owner_id=_required(body, "owner_id"),
            title=_required(body, "title"),
            description=str(body.get("description", "")),
            evidence=body.get("evidence") or [],
            skills=body.get("skills") or [],
            published=bool(body.get("published", False)),
        )
        session.add(item)
        session.flush()
        _audit(session, item.owner_id, "create", "portfolio_item", item.id)
        return _as_dict(item, "id", "title", "description", "evidence", "skills", "published")


@router.get("/cpd/{user_id}")
def list_cpd(user_id: str):
    with session_scope() as session:
        rows = session.scalars(select(CPDRecord).where(CPDRecord.owner_id == user_id)).all()
        return [_as_dict(row, "id", "title", "provider", "completed_on", "hours", "evidence") for row in rows]


@router.post("/cpd")
def create_cpd(body: dict):
    with session_scope() as session:
        record = CPDRecord(
            owner_id=_required(body, "owner_id"),
            title=_required(body, "title"),
            provider=body.get("provider"),
            completed_on=body.get("completed_on"),
            hours=float(body.get("hours") or 0),
            evidence=body.get("evidence") or {},
        )
        session.add(record)
        session.flush()
        _audit(session, record.owner_id, "create", "cpd_record", record.id)
        return _as_dict(record, "id", "title", "provider", "completed_on", "hours", "evidence")


@router.get("/career/pathways")
def career_pathways():
    return {
        "pathways": [
            {"id": "software", "title": "Software & AI", "skills": ["Programming", "Algorithms", "AI", "Cloud"], "projects": ["Portfolio app", "ML case study"]},
            {"id": "data", "title": "Data & Analytics", "skills": ["Statistics", "SQL", "Visualisation", "ML"], "projects": ["Data report", "Predictive model"]},
            {"id": "business", "title": "Business & Leadership", "skills": ["Strategy", "Finance", "Communication", "Operations"], "projects": ["Business case", "Financial model"]},
            {"id": "research", "title": "Academic Research", "skills": ["Methods", "Literature review", "Analysis", "Writing"], "projects": ["Research proposal", "Capstone"]},
        ]
    }


# Institutions, cohorts, discussions and LMS ---------------------------------

@router.get("/courses")
def list_courses(level_id: str = "", subject: str = "", limit: int = 200):
    with session_scope() as session:
        statement = select(Course).where(Course.deleted_at.is_(None))
        if level_id:
            statement = statement.where(Course.level_id == level_id)
        if subject:
            statement = statement.where(Course.subject == subject)
        rows = session.scalars(statement.limit(max(1, min(limit, 1000)))).all()
        return [
            _as_dict(row, "id", "level_id", "subject", "title", "description", "published")
            for row in rows
        ]


@router.get("/organizations")
def list_organizations():
    with session_scope() as session:
        rows = session.scalars(select(Organization).where(Organization.deleted_at.is_(None))).all()
        return [_as_dict(row, "id", "name", "kind", "settings") for row in rows]


@router.post("/organizations")
def create_organization(body: dict):
    with session_scope() as session:
        organization = Organization(
            name=_required(body, "name"),
            kind=str(body.get("kind") or "university"),
            settings=body.get("settings") or {},
        )
        session.add(organization)
        session.flush()
        owner_id = body.get("owner_id")
        if owner_id:
            session.add(Membership(organization_id=organization.id, user_id=owner_id, role="admin"))
        _audit(session, owner_id, "create", "organization", organization.id, {"name": organization.name})
        return _as_dict(organization, "id", "name", "kind", "settings")


@router.post("/cohorts")
def create_cohort(body: dict):
    with session_scope() as session:
        cohort = Cohort(
            organization_id=_required(body, "organization_id"),
            course_id=body.get("course_id") or None,
            name=_required(body, "name"),
            settings=body.get("settings") or {},
        )
        session.add(cohort)
        session.flush()
        return _as_dict(cohort, "id", "organization_id", "course_id", "name", "settings")


@router.get("/cohorts/{cohort_id}/attendance")
def list_attendance(cohort_id: str):
    with session_scope() as session:
        rows = session.scalars(
            select(AttendanceRecord)
            .where(AttendanceRecord.cohort_id == cohort_id, AttendanceRecord.deleted_at.is_(None))
            .order_by(AttendanceRecord.session_date.desc())
        ).all()
        return [
            _as_dict(row, "id", "cohort_id", "user_id", "session_date", "status", "note")
            for row in rows
        ]


@router.post("/cohorts/{cohort_id}/attendance")
def record_attendance(cohort_id: str, body: dict):
    user_id = _required(body, "user_id")
    session_date = _required(body, "session_date")
    with session_scope() as session:
        row = session.scalar(
            select(AttendanceRecord).where(
                AttendanceRecord.cohort_id == cohort_id,
                AttendanceRecord.user_id == user_id,
                AttendanceRecord.session_date == session_date,
            )
        )
        before = None
        if row:
            before = {"status": row.status, "note": row.note}
            row.status = str(body.get("status") or "present")
            row.note = str(body.get("note") or "")
        else:
            row = AttendanceRecord(
                cohort_id=cohort_id,
                user_id=user_id,
                session_date=session_date,
                status=str(body.get("status") or "present"),
                note=str(body.get("note") or ""),
            )
            session.add(row)
            session.flush()
        result = _as_dict(row, "id", "cohort_id", "user_id", "session_date", "status", "note")
        _audit(session, user_id, "attendance_recorded", "attendance", row.id, result, before=before)
        return result


@router.get("/discussions/{course_id}")
def list_discussions(course_id: str):
    with session_scope() as session:
        rows = session.scalars(select(DiscussionPost).where(DiscussionPost.course_id == course_id)).all()
        return [_as_dict(row, "id", "course_id", "author_id", "parent_id", "body", "created_at") for row in rows]


@router.post("/discussions")
def create_discussion(body: dict):
    with session_scope() as session:
        post = DiscussionPost(
            course_id=_required(body, "course_id"),
            author_id=_required(body, "author_id"),
            parent_id=body.get("parent_id") or None,
            body=_required(body, "body"),
        )
        session.add(post)
        session.flush()
        return _as_dict(post, "id", "course_id", "author_id", "parent_id", "body", "created_at")


@router.get("/lms/scorm/{course_id}")
def export_scorm(course_id: str):
    with session_scope() as session:
        course = session.get(Course, course_id)
        if not course:
            raise HTTPException(404, "Course not found")
        payload = _as_dict(course, "id", "level_id", "subject", "title", "description")
    manifest = f"""<?xml version="1.0" encoding="UTF-8"?>
<manifest identifier="{course_id}" version="1.2"
 xmlns="http://www.imsproject.org/xsd/imscp_rootv1p1p2">
 <organizations default="eduai"><organization identifier="eduai">
  <title>{payload['title']}</title>
  <item identifier="course" identifierref="resource"><title>{payload['title']}</title></item>
 </organization></organizations>
 <resources><resource identifier="resource" type="webcontent" href="course.json">
  <file href="course.json"/>
 </resource></resources>
</manifest>"""
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("imsmanifest.xml", manifest)
        archive.writestr("course.json", json.dumps(payload, indent=2))
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{course_id}-scorm.zip"'},
    )


@router.post("/lms/xapi")
def xapi_statement(body: dict):
    return {
        "id": str(uuid.uuid4()),
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "actor": body.get("actor"),
        "verb": body.get("verb"),
        "object": body.get("object"),
        "result": body.get("result") or {},
        "stored": False,
        "note": "Set an external LRS endpoint and credentials to transmit this xAPI statement.",
    }


@router.get("/lms/lti/config")
def lti_config():
    return {
        "title": "EduAI_Pro",
        "description": "Adult learning, research and professional development",
        "oidc_initiation_url": "/api/pro/lms/lti/login",
        "target_link_uri": "/api/pro/lms/lti/launch",
        "jwks_url": "/api/pro/lms/lti/jwks",
        "scopes": [
            "https://purl.imsglobal.org/spec/lti-ags/scope/score",
            "https://purl.imsglobal.org/spec/lti-nrps/scope/contextmembership.readonly",
        ],
        "status": "configuration scaffold; register platform keys before production launch",
    }


# ─── Resume Builder ─────────────────────────────────────────────────────────
# Turns a candidate's raw skills/portfolio notes and an optional target job
# description into a polished, quantified resume draft, downloadable as PDF
# or DOCX. Fills the "CV & cover letter" capability the Career Lab UI has
# advertised without an implementation behind it.

@router.post("/resume/draft")
def draft_resume(body: dict):
    skills = [str(s).strip() for s in (body.get("skills") or []) if str(s).strip()][:30]
    experience = [str(s).strip() for s in (body.get("experience") or []) if str(s).strip()][:20]
    if not skills and not experience:
        raise HTTPException(400, "Add at least one skill or experience note before drafting a resume")

    profile = {
        "name": str(body.get("name", "")).strip(),
        "target_role": str(body.get("target_role", "")).strip(),
        "job_description": str(body.get("job_description", ""))[:4000],
        "skills": skills,
        "experience": experience,
        "education": str(body.get("education", ""))[:1000],
    }
    return ai_tutor.generate_resume_content(profile)


def _resume_wrapped_lines(text: str, width: int = 95) -> list[str]:
    lines = []
    for paragraph in text.splitlines() or [text]:
        lines.extend(textwrap.wrap(paragraph, width) or [""])
    return lines


def _resume_pdf(name: str, target_role: str, contact: str, education: str, draft: dict) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 72

    def new_page():
        nonlocal y
        pdf.showPage()
        y = height - 72

    def heading(text: str, size: int = 13):
        nonlocal y
        if y < 100:
            new_page()
        pdf.setFont("Helvetica-Bold", size)
        pdf.drawString(72, y, text)
        y -= 18

    def body(text: str, indent: int = 72):
        nonlocal y
        pdf.setFont("Helvetica", 10)
        for line in _resume_wrapped_lines(text):
            if y < 72:
                new_page()
                pdf.setFont("Helvetica", 10)
            pdf.drawString(indent, y, line)
            y -= 14

    heading(name or "Resume", size=18)
    if target_role:
        body(target_role)
    if contact:
        body(contact)
    y -= 10

    if draft.get("summary"):
        heading("Summary")
        body(draft["summary"])
        y -= 8
    if draft.get("skills"):
        heading("Skills")
        body(", ".join(draft["skills"]))
        y -= 8
    if draft.get("experience"):
        heading("Experience")
        for bullet in draft["experience"]:
            body(f"- {bullet}")
        y -= 8
    if education:
        heading("Education")
        body(education)

    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _resume_docx(name: str, target_role: str, contact: str, education: str, draft: dict) -> bytes:
    doc = Document()
    doc.add_heading(name or "Resume", level=1)
    if target_role:
        doc.add_paragraph(target_role)
    if contact:
        doc.add_paragraph(contact)
    if draft.get("summary"):
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(draft["summary"])
    if draft.get("skills"):
        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(draft["skills"]))
    if draft.get("experience"):
        doc.add_heading("Experience", level=2)
        for bullet in draft["experience"]:
            doc.add_paragraph(bullet, style="List Bullet")
    if education:
        doc.add_heading("Education", level=2)
        doc.add_paragraph(education)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@router.post("/resume/export")
def export_resume(body: dict):
    draft = {
        "summary": str(body.get("summary", "")).strip(),
        "skills": [str(s).strip() for s in (body.get("skills") or []) if str(s).strip()],
        "experience": [str(s).strip() for s in (body.get("experience") or []) if str(s).strip()],
    }
    if not draft["summary"] and not draft["skills"] and not draft["experience"]:
        raise HTTPException(400, "Nothing to export -- draft a resume first")

    name = str(body.get("name", "")).strip()
    target_role = str(body.get("target_role", "")).strip()
    contact = str(body.get("contact", "")).strip()
    education = str(body.get("education", "")).strip()
    format_ = body.get("format", "pdf")

    filename_base = re.sub(r"[^a-zA-Z0-9-]+", "-", name or "resume").strip("-").lower() or "resume"
    if format_ == "pdf":
        return StreamingResponse(
            BytesIO(_resume_pdf(name, target_role, contact, education, draft)),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.pdf"'},
        )
    if format_ == "docx":
        return StreamingResponse(
            BytesIO(_resume_docx(name, target_role, contact, education, draft)),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename_base}.docx"'},
        )
    raise HTTPException(422, "format must be 'pdf' or 'docx'")
