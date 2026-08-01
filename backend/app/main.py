import hashlib
import json
import os
import sqlite3 as _sqlite3
import subprocess
import tempfile
import time
from contextlib import asynccontextmanager
from functools import lru_cache
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import quote, quote_plus
from urllib.request import Request, urlopen

from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from sqlalchemy import select

from app.database import create_schema, session_scope
from app.models import Resource
from app.safety import safety_filter
from app.storage import (
    ALLOWED_CHILDREN,
    ALL_PROFILES,
    get_children,
    get_all_profiles,
    get_all_parent_profiles,
    add_user,
    rename_user,
    delete_user,
    get_progress,
    save_progress,
    delete_snippet,
    get_activity_log,
    append_activity,
    get_homework,
    save_homework,
    get_reading_log,
    append_reading_entry,
    get_screen_time,
    add_screen_time,
    get_attendance as get_attendance_records,
    save_attendance as save_attendance_records,
)
from app.websearch import web_search, SearchNotConfigured
from app.curate import curate_resource, CurationError, RESOURCE_KEYS as CURATE_RESOURCE_KEYS
from app.summarize import summarize
from app import resource_tab
from app import paintings as paintings_store
from app import ai_tutor
from app import content_store
from app import levels as levels_module
from app import course_catalog
from app import local_library
from app import ai_reliability
from app import personalized_learning
from app import pdf_explainer
from app import lesson_planner
from app.professional_api import router as professional_router

BASE_DIR = Path(__file__).resolve().parent.parent
SYLLABUS_DIR = BASE_DIR / "syllabus"
SAFE_DIR = BASE_DIR / "safe"

@asynccontextmanager
async def lifespan(_app: FastAPI):
    """Make a fresh local install usable before its first dashboard request."""
    # Production deployments still use Alembic migrations. This idempotent
    # bootstrap covers portable/local SQLite installs whose database has not
    # yet been created.
    create_schema()
    yield


app = FastAPI(
    title="Global Education Platform API",
    description="An all-ages learning platform spanning school, college, undergraduate, and master's levels.",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)
app.include_router(professional_router)

_museum_resource_dir = Path(__file__).parent.parent / "data" / "museum_resource"
if _museum_resource_dir.exists():
    app.mount("/museum-resource", StaticFiles(directory=str(_museum_resource_dir)), name="museum-resource")

_movie_thumbnail_dir = Path(__file__).parent.parent / "data" / "movie_thumbnails"
if _movie_thumbnail_dir.exists():
    app.mount("/movie-thumbnails", StaticFiles(directory=str(_movie_thumbnail_dir)), name="movie-thumbnails")


def _require_child(child: str) -> str:
    if child not in get_children():
        raise HTTPException(status_code=404, detail="Unknown child profile")
    return child


def _sanitize_json(obj, strict: bool = True):
    if isinstance(obj, dict):
        return {k: _sanitize_json(v, strict=strict) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_sanitize_json(v, strict=strict) for v in obj]
    if isinstance(obj, str):
        return safety_filter.sanitize(obj, strict=strict)
    return obj


# Grade/level syllabus files are large (tens of MB at the college/master's
# levels after the full lesson-count expansion) and are read-heavy -- every
# request to /api/grade/{n} or /api/level/{id} used to re-parse the whole
# file and re-run the recursive safety sanitizer over every string in it from
# scratch. At this data volume that's expensive enough to make a test suite
# (or a few concurrent users) that touches these endpoints repeatedly grind
# to a crawl. Cache the parsed+sanitized result keyed by the file's mtime, so
# a rewrite (e.g. from /api/apply-link-fixes or a content-generation script)
# still naturally busts the cache on next read.
@lru_cache(maxsize=64)
def _load_syllabus_json(path_str: str, mtime: float) -> dict:
    with open(path_str, encoding="utf-8") as f:
        return json.load(f)


@lru_cache(maxsize=128)
def _load_sanitized_syllabus(path_str: str, mtime: float, strict: bool) -> dict:
    return _sanitize_json(_load_syllabus_json(path_str, mtime), strict=strict)


@app.get("/api/grade/{standard}")
def get_grade(standard: int):
    path = SYLLABUS_DIR / f"grade{standard}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grade {standard} not available yet")
    data = json.loads(json.dumps(_load_sanitized_syllabus(str(path), path.stat().st_mtime, True)))
    with session_scope() as session:
        curated = list(
            session.scalars(
                select(Resource).where(
                    Resource.metadata_json["runtime_curated"].as_boolean().is_(True),
                    Resource.metadata_json["level_id"].as_string() == str(standard),
                    Resource.deleted_at.is_(None),
                )
            )
        )
    for row in curated:
        metadata = row.metadata_json or {}
        subject = metadata.get("subject")
        resource_type = metadata.get("resource_type")
        if not subject or resource_type not in CURATE_RESOURCE_KEYS:
            continue
        subject_data = data.setdefault("subjects", {}).setdefault(subject, {})
        subject_data.setdefault(resource_type, []).append(metadata.get("resource") or {"title": row.title, "url": row.url})
    return data


@app.get("/api/progress/{child}")
def read_progress(child: str):
    _require_child(child)
    return get_progress(child)


@app.post("/api/progress/{child}")
def update_progress(child: str, update: dict):
    _require_child(child)
    result = save_progress(child, update)
    append_activity(child, {"type": "progress_update", "data": update})
    return result


@app.delete("/api/progress/{child}/snippets/{snippet_id}")
def remove_progress_snippet(child: str, snippet_id: str):
    _require_child(child)
    return delete_snippet(child, snippet_id)


def _progress_csv(child: str, progress: dict) -> str:
    lines = ["subject,score"]
    for subject, score in progress.get("scores", {}).items():
        lines.append(f"{subject},{score}")
    lines.append("")
    lines.append(f"badges,{';'.join(progress.get('badges', []))}")
    return "\n".join(lines)


def _progress_pdf(child: str, progress: dict) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 72
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, y, f"{child}'s Progress Report")
    y -= 30
    pdf.setFont("Helvetica", 12)
    for subject, score in progress.get("scores", {}).items():
        pdf.drawString(72, y, f"{subject}: {score}")
        y -= 18
    y -= 12
    pdf.drawString(72, y, f"Badges: {', '.join(progress.get('badges', [])) or 'None yet'}")
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


@app.get("/api/progress/{child}/export")
def export_progress(child: str, format: str = "csv"):
    _require_child(child)
    progress = get_progress(child)
    if format == "csv":
        return StreamingResponse(
            BytesIO(_progress_csv(child, progress).encode("utf-8")),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{child}-progress.csv"'},
        )
    if format == "pdf":
        return StreamingResponse(
            BytesIO(_progress_pdf(child, progress)),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{child}-progress.pdf"'},
        )
    raise HTTPException(status_code=422, detail="format must be 'csv' or 'pdf'")


@app.get("/api/grade/{standard}/export")
def export_syllabus(standard: int, format: str = "json"):
    path = SYLLABUS_DIR / f"grade{standard}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grade {standard} not available yet")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    data = _sanitize_json(data)

    if format == "json":
        return StreamingResponse(
            BytesIO(json.dumps(data, indent=2).encode("utf-8")),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="grade{standard}-syllabus.json"'},
        )
    if format == "csv":
        lines = ["subject,resource_type,title,url"]
        for subject, content in data.get("subjects", {}).items():
            for resource_type, items in content.items():
                if not isinstance(items, list):
                    continue
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    title = item.get("title", "")
                    url = item.get("url") or item.get("link", "")
                    lines.append(f'"{subject}","{resource_type}","{title}","{url}"')
        return StreamingResponse(
            BytesIO("\n".join(lines).encode("utf-8")),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="grade{standard}-syllabus.csv"'},
        )
    raise HTTPException(status_code=422, detail="format must be 'json' or 'csv'")


def _filtered_subjects(data: dict, subjects: list, resource_types: list) -> dict:
    all_subjects = data.get("subjects", {})
    chosen_subjects = subjects or list(all_subjects.keys())
    result = {}
    for subject_name in chosen_subjects:
        content = all_subjects.get(subject_name)
        if not content:
            continue
        chosen_types = resource_types or list(content.keys())
        result[subject_name] = {k: content[k] for k in chosen_types if k in content}
    return result


def _custom_docx(standard: int, filtered: dict) -> bytes:
    doc = Document()
    doc.add_heading(f"Grade {standard} Syllabus", level=1)
    for subject_name, content in filtered.items():
        doc.add_heading(subject_name, level=2)
        for resource_type, items in content.items():
            if not isinstance(items, list) or not items:
                continue
            doc.add_heading(resource_type.replace("_", " ").title(), level=3)
            for item in items:
                if not isinstance(item, dict):
                    continue
                title = item.get("title", "")
                url = item.get("link") or item.get("url", "")
                fact = item.get("fact", "")
                line = title
                if url:
                    line += f" — {url}"
                if fact:
                    line += f": {fact}"
                doc.add_paragraph(line, style="List Bullet")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _custom_pdf(standard: int, filtered: dict) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 72

    def new_page():
        nonlocal y
        pdf.showPage()
        y = height - 72

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, y, f"Grade {standard} Syllabus")
    y -= 30
    for subject_name, content in filtered.items():
        if y < 100:
            new_page()
        pdf.setFont("Helvetica-Bold", 13)
        pdf.drawString(72, y, subject_name)
        y -= 20
        for resource_type, items in content.items():
            if not isinstance(items, list) or not items:
                continue
            if y < 100:
                new_page()
            pdf.setFont("Helvetica-Oblique", 11)
            pdf.drawString(90, y, resource_type.replace("_", " ").title())
            y -= 16
            pdf.setFont("Helvetica", 10)
            for item in items:
                if not isinstance(item, dict):
                    continue
                title = str(item.get("title", ""))[:90]
                if y < 80:
                    new_page()
                pdf.drawString(108, y, f"- {title}")
                y -= 14
        y -= 8
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


@app.post("/api/grade/{standard}/export/custom")
def export_syllabus_custom(standard: int, payload: dict):
    path = SYLLABUS_DIR / f"grade{standard}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grade {standard} not available yet")
    data = _load_sanitized_syllabus(str(path), path.stat().st_mtime, True)

    subjects = payload.get("subjects") or []
    resource_types = payload.get("resource_types") or []
    format_ = payload.get("format", "pdf")

    filtered = _filtered_subjects(data, subjects, resource_types)

    if format_ == "pdf":
        return StreamingResponse(
            BytesIO(_custom_pdf(standard, filtered)),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="grade{standard}-custom.pdf"'},
        )
    if format_ == "docx":
        return StreamingResponse(
            BytesIO(_custom_docx(standard, filtered)),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="grade{standard}-custom.docx"'},
        )
    raise HTTPException(status_code=422, detail="format must be 'pdf' or 'docx'")


# ─── All-ages Levels (school, college, undergraduate, master's) ────────────
# Generalizes /api/grade/{standard} to every level the platform understands,
# including the new C1, C2, UG1-UG4, M1, M2 levels. The backend is the single
# source of truth for which levels exist (see app/levels.py) so the frontend
# never has to hardcode the level list.

def _level_syllabus_path(level_id: str) -> Path:
    return SYLLABUS_DIR / levels_module.syllabus_filename(level_id)


@app.get("/api/levels")
def list_levels():
    return {"levels": levels_module.all_levels()}


def _level_path_and_mode(level_id: str) -> tuple[Path, bool, str]:
    if not levels_module.is_valid_level(level_id):
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' is not recognised")
    path = _level_syllabus_path(level_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' not available yet")
    return path, levels_module.is_school_level(level_id), levels_module.normalize_level_id(level_id)


@lru_cache(maxsize=1024)
def _load_sanitized_subject(path_str: str, mtime: float, strict: bool, subject_name: str) -> dict:
    raw = _load_syllabus_json(path_str, mtime).get("subjects", {}).get(subject_name)
    if raw is None:
        raise KeyError(subject_name)
    return _sanitize_json(raw, strict=strict)


def _trusted_video_link(subject_name: str, lesson_title: str) -> dict:
    academic = any(
        token in subject_name.lower()
        for token in ("engineering", "computer", "coding", "physics", "chemistry", "math", "data", "machine", "artificial")
    )
    source = "MIT OpenCourseWare" if academic else "Khan Academy"
    query = quote_plus(f"{source} {subject_name} {lesson_title}")
    return {
        "title": f"{lesson_title} — trusted video lesson",
        "url": f"https://www.youtube.com/results?search_query={query}",
        "provider": source,
        "description": f"Search {source}'s educational videos for this exact lesson.",
        "safe": True,
    }


def _technical_enrichment(subject_name: str, lesson: dict) -> dict:
    """Add compact, deterministic teaching aids without inflating source JSON files."""
    title = str(lesson.get("title") or "this topic")
    key_concepts = [str(item) for item in lesson.get("key_concepts", [])[:4]]
    concept = key_concepts[0] if key_concepts else title
    lower = subject_name.lower()
    quantitative = any(
        token in lower
        for token in ("math", "physics", "chemistry", "engineering", "economics", "finance", "statistics", "data", "computer")
    )
    if quantitative:
        formula = lesson.get("formula") or "result = known inputs × applicable rate or relationship"
        worked = [
            "Identify the known quantities, units and required result.",
            f"Select the governing relationship for {concept}: {formula}.",
            "Substitute values, calculate carefully, then check units and scale.",
            "Interpret the result in the original real-world context.",
        ]
        graph = {
            "title": f"How the main variables in {title} relate",
            "x_axis": "Independent variable",
            "y_axis": "Measured outcome",
            "points": [0, 2, 3, 5, 8, 13],
        }
    else:
        formula = "Claim + relevant evidence + reasoning = defensible conclusion"
        worked = [
            "Define the question and important terms.",
            "Collect one primary and one reliable secondary source.",
            "Compare evidence, assumptions and competing interpretations.",
            "State a qualified conclusion and identify its limitations.",
        ]
        graph = {
            "title": f"Evidence strength across a {title} investigation",
            "x_axis": "Investigation stage",
            "y_axis": "Evidence strength",
            "points": [1, 2, 4, 5, 7, 8],
        }
    return {
        "technical_detail": lesson.get("technical_detail")
        or f"{title} is analysed through precise definitions, assumptions, mechanisms, evidence and limitations. "
           f"Track how {concept} changes when one condition varies while the others are controlled.",
        "formulae": lesson.get("formulae") or [formula],
        "worked_example": lesson.get("worked_example") or {
            "problem": f"Apply {title} to a realistic decision with incomplete information.",
            "steps": worked,
            "answer": "A sound answer shows the method, checks evidence or units, and explains what the result means.",
        },
        "real_world_example": lesson.get("real_world_example")
        or f"Professionals use {title} to compare alternatives, justify decisions and communicate risk in real projects.",
        "practical_problem": lesson.get("practical_problem")
        or f"Choose a local or workplace example of {title}; record inputs or evidence, apply the method, and defend your conclusion.",
        "data_table": lesson.get("data_table") or {
            "headers": ["Stage", "Input or evidence", "Method", "Check"],
            "rows": [
                ["1", "Baseline information", "Define and classify", "Is the source reliable?"],
                ["2", "Observed change", "Calculate or compare", "Are units/terms consistent?"],
                ["3", "Result", "Interpret and explain", "Does it answer the question?"],
            ],
        },
        "graph": lesson.get("graph") or graph,
        "figure": lesson.get("figure") or {
            "caption": f"Concept map for {title}",
            "nodes": ["Inputs", concept, "Method", "Result", "Evaluation"],
        },
        "video_resources": lesson.get("video_resources") or [_trusted_video_link(subject_name, title)],
    }


@app.get("/api/level/{level_id}/overview")
def get_level_overview(level_id: str):
    path, _strict, norm = _level_path_and_mode(level_id)
    data = _load_syllabus_json(str(path), path.stat().st_mtime)
    subjects = {}
    for name, subject in data.get("subjects", {}).items():
        subjects[name] = {
            "lesson_count": len(subject.get("lessons", [])),
            "resource_count": sum(
                len(subject.get(key, []))
                for key in ("books", "textbooks", "text_resources", "video_resources", "external_courses")
            ),
        }
    return {
        "level": norm,
        "standard": data.get("standard"),
        "level_info": levels_module.get_level(level_id),
        "subjects": subjects,
    }


@app.get("/api/level/{level_id}/subjects/{subject_name}")
def get_level_subject(level_id: str, subject_name: str):
    path, strict, norm = _level_path_and_mode(level_id)
    try:
        cached = _load_sanitized_subject(str(path), path.stat().st_mtime, strict, subject_name)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail=f"Subject '{subject_name}' is not available at level {norm}") from exc
    subject = json.loads(json.dumps(cached))
    subject["lessons"] = [
        {**lesson, **_technical_enrichment(subject_name, lesson)}
        for lesson in subject.get("lessons", [])
    ]
    subject["__name"] = subject_name
    return {"level": norm, "subject_name": subject_name, "subject": subject}


class LearningEvidence(BaseModel):
    level_id: str
    subject: str
    concept: str
    correct: bool
    lesson_id: str = ""
    question_id: str = ""
    answer: str = ""
    expected_answer: str = ""
    confidence: float = 1.0


def _adaptive_subject(level_id: str, subject_name: str) -> tuple[str, dict]:
    path, _strict, normalized = _level_path_and_mode(level_id)
    data = _load_syllabus_json(str(path), path.stat().st_mtime)
    subject = data.get("subjects", {}).get(subject_name)
    if subject is None:
        raise HTTPException(
            status_code=404,
            detail=f"Subject '{subject_name}' is not available at level {normalized}",
        )
    return normalized, subject


@app.get("/api/personalized/{profile}/{level_id}/{subject_name}")
def personalized_profile(profile: str, level_id: str, subject_name: str):
    normalized, subject = _adaptive_subject(level_id, subject_name)
    try:
        return personalized_learning.build_profile(profile, normalized, subject_name, subject)
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.post("/api/personalized/{profile}/evidence")
def personalized_evidence(profile: str, body: LearningEvidence):
    _adaptive_subject(body.level_id, body.subject)
    try:
        return personalized_learning.record_evidence(
            profile,
            levels_module.normalize_level_id(body.level_id),
            body.subject,
            body.concept,
            body.correct,
            lesson_id=body.lesson_id,
            question_id=body.question_id,
            answer=body.answer,
            expected_answer=body.expected_answer,
            confidence=body.confidence,
        )
    except LookupError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc


@app.get("/api/level/{level_id}")
def get_level_content(level_id: str):
    path, strict, normalized = _level_path_and_mode(level_id)
    data = dict(_load_sanitized_syllabus(str(path), path.stat().st_mtime, strict))
    data.setdefault("level", normalized)
    data["level_info"] = levels_module.get_level(level_id)
    return data


@app.get("/api/level/{level_id}/search")
def search_level(level_id: str, q: str):
    if not levels_module.is_valid_level(level_id):
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' is not recognised")
    path = _level_syllabus_path(level_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' not available yet")
    data = _load_syllabus_json(str(path), path.stat().st_mtime)

    query = q.strip().lower()
    if not query:
        return []

    strict = levels_module.is_school_level(level_id)
    results = []
    for subject_name, subject in data.get("subjects", {}).items():
        for key in RESOURCE_KEYS:
            for resource in subject.get(key, []):
                if not isinstance(resource, dict) or resource.get("safe") is not True:
                    continue
                haystack = " ".join(
                    str(v) for v in (resource.get("title"), resource.get("description"))
                    if v
                ).lower()
                if query in haystack:
                    results.append(
                        _sanitize_json({**resource, "subject": subject_name, "resource_type": key}, strict=strict)
                    )
    return results


@app.get("/api/level/{level_id}/export")
def export_level(level_id: str, format: str = "json"):
    if not levels_module.is_valid_level(level_id):
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' is not recognised")
    path = _level_syllabus_path(level_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Level '{level_id}' not available yet")
    strict = levels_module.is_school_level(level_id)
    data = _load_sanitized_syllabus(str(path), path.stat().st_mtime, strict)
    norm = levels_module.normalize_level_id(level_id)

    if format == "json":
        return StreamingResponse(
            BytesIO(json.dumps(data, indent=2).encode("utf-8")),
            media_type="application/json",
            headers={"Content-Disposition": f'attachment; filename="level-{norm}-syllabus.json"'},
        )
    if format == "csv":
        lines = ["subject,resource_type,title,url"]
        for subject, content in data.get("subjects", {}).items():
            for resource_type, items in content.items():
                if not isinstance(items, list):
                    continue
                for item in items:
                    if not isinstance(item, dict):
                        continue
                    title = item.get("title", "")
                    url = item.get("url") or item.get("link", "")
                    lines.append(f'"{subject}","{resource_type}","{title}","{url}"')
        return StreamingResponse(
            BytesIO("\n".join(lines).encode("utf-8")),
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="level-{norm}-syllabus.csv"'},
        )
    raise HTTPException(status_code=422, detail="format must be 'json' or 'csv'")


# ─── Music & Instruments ─────────────────────────────────────────────────────
_MUSIC_INSTRUMENTS_PATH = Path(__file__).parent.parent / "data" / "music_instruments" / "music.json"


@lru_cache(maxsize=1)
def _load_music_instruments() -> dict:
    with open(_MUSIC_INSTRUMENTS_PATH, encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/music-instruments")
def music_instruments_overview():
    data = _load_music_instruments()
    return {
        "title": data["title"],
        "description": data["description"],
        "categories": [
            {"id": k, "label": v["label"], "emoji": v["emoji"], "description": v["description"]}
            for k, v in data["categories"].items()
        ],
        "instruments": [
            {"id": k, "label": v["label"], "emoji": v["emoji"]}
            for k, v in data["instruments"].items()
        ],
    }


@app.get("/api/music-instruments/category/{category_id}")
def music_instruments_category(category_id: str):
    data = _load_music_instruments()
    category = data["categories"].get(category_id)
    if not category:
        raise HTTPException(status_code=404, detail="Music category not found")
    return {"id": category_id, **category}


@app.get("/api/music-instruments/instrument/{instrument_id}")
def music_instruments_instrument(instrument_id: str):
    data = _load_music_instruments()
    instrument = data["instruments"].get(instrument_id)
    if not instrument:
        raise HTTPException(status_code=404, detail="Instrument not found")
    return {"id": instrument_id, **instrument}


@app.post("/api/exam-result/export")
def export_exam_result(payload: dict):
    try:
        child = payload["child"]
        subject = payload["subject"]
        score = payload["score"]
        passed = payload["passed"]
        answers = payload.get("answers", [])
    except KeyError as exc:
        raise HTTPException(status_code=422, detail="child, subject, score, and passed are required") from exc

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 72
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(72, y, f"{child}'s Exam Result: {subject}")
    y -= 30
    pdf.setFont("Helvetica", 12)
    pdf.drawString(72, y, f"Score: {score}% ({'Passed' if passed else 'Not yet passed'})")
    y -= 24
    for i, answer in enumerate(answers, start=1):
        question = str(answer.get("question", ""))[:90]
        given = str(answer.get("given", ""))[:60]
        pdf.drawString(72, y, f"{i}. {question}")
        y -= 16
        pdf.drawString(90, y, f"Answer: {given}")
        y -= 20
        if y < 72:
            pdf.showPage()
            y = height - 72
    pdf.showPage()
    pdf.save()

    return StreamingResponse(
        BytesIO(buffer.getvalue()),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{child}-{subject}-exam-result.pdf"'},
    )


ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".txt", ".png", ".jpg", ".jpeg", ".mp3", ".wav"}


def _content_type_error(ext: str, contents: bytes) -> str | None:
    if ext == ".png" and not contents.startswith(b"\x89PNG\r\n\x1a\n"):
        return "PNG upload rejected: file signature does not match .png"
    if ext in {".jpg", ".jpeg"} and not contents.startswith(b"\xff\xd8\xff"):
        return "JPEG upload rejected: file signature does not match .jpg/.jpeg"
    if ext == ".wav" and not (contents.startswith(b"RIFF") and contents[8:12] == b"WAVE"):
        return "WAV upload rejected: file signature does not match .wav"
    if ext == ".mp3":
        has_id3 = contents.startswith(b"ID3")
        has_frame = len(contents) >= 2 and contents[0] == 0xFF and (contents[1] & 0xE0) == 0xE0
        if not (has_id3 or has_frame):
            return "MP3 upload rejected: file signature does not match .mp3"
    if ext == ".pdf" and contents and not contents.lstrip().startswith(b"%PDF"):
        return "PDF upload rejected: file signature does not match .pdf"
    return None


@app.post("/api/upload-safe-book")
async def upload_safe_book(
    file: UploadFile = File(...),
    standard: int | None = Form(None),
    subject: str | None = Form(None),
):
    filename = file.filename or ""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed: {sorted(ALLOWED_UPLOAD_EXTENSIONS)}",
        )
    if not safety_filter.is_safe(filename):
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    contents = await file.read()
    type_error = _content_type_error(ext, contents)
    if type_error:
        raise HTTPException(status_code=400, detail=type_error)
    full_text = ""
    if ext == ".txt":
        full_text = contents.decode("utf-8", errors="ignore")
    elif ext == ".pdf":
        try:
            reader = PdfReader(BytesIO(contents))
            full_text = "".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            full_text = ""

    text_sample = full_text[:5000]
    if text_sample and not safety_filter.is_safe(text_sample):
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    summary = summarize(full_text) if full_text else ""
    if summary and not safety_filter.is_safe(summary):
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    result = {"filename": filename, "status": "accepted", "type": ext.lstrip("."), "summary": summary}

    if standard is not None and subject:
        if not summary:
            raise HTTPException(
                status_code=400,
                detail="Could not extract readable text from this file to summarize and add it.",
            )
        try:
            saved = curate_resource(
                standard,
                subject,
                "text_resources",
                {"title": Path(filename).stem, "description": summary},
            )
        except CurationError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
        result["added_resource"] = saved

    return result


@app.post("/api/resource-tab/upload")
async def resource_tab_upload(file: UploadFile = File(...)):
    filename = file.filename or ""
    if not safety_filter.is_safe(filename):
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    contents = await file.read()
    try:
        record = resource_tab.add_document(filename, contents)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if record["summary"] and not safety_filter.is_safe(record["summary"]):
        resource_tab.delete_document(record["id"])
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    return record


@app.get("/api/resource-tab")
def resource_tab_list():
    return resource_tab.list_documents()


@app.get("/api/resource-tab/{doc_id}/download")
def resource_tab_download(doc_id: str):
    record = resource_tab.get_document(doc_id)
    path = resource_tab.get_document_path(doc_id)
    if not record or not path or not path.exists():
        raise HTTPException(status_code=404, detail="Document not found")
    media_types = {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "txt": "text/plain",
        "md": "text/markdown",
        "rtf": "application/rtf",
        "html": "text/html",
        "htm": "text/html",
        "epub": "application/epub+zip",
        "mobi": "application/x-mobipocket-ebook",
        "azw": "application/vnd.amazon.ebook",
        "azw3": "application/vnd.amazon.ebook",
        "kfx": "application/vnd.amazon.ebook",
        "fb2": "application/x-fictionbook+xml",
        "odt": "application/vnd.oasis.opendocument.text",
    }
    with open(path, "rb") as f:
        data = f.read()
    return StreamingResponse(
        BytesIO(data),
        media_type=media_types.get(record["type"], "application/octet-stream"),
        headers={"Content-Disposition": f'attachment; filename="{record["filename"]}"'},
    )


@app.delete("/api/resource-tab/{doc_id}")
def resource_tab_delete(doc_id: str):
    if not resource_tab.delete_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}


# ─── PDF Explainer ───────────────────────────────────────────────────────────
# Upload any PDF and get an AI-simplified explanation (readable aloud via the
# browser's built-in text-to-speech, like other content in this app), ask
# follow-up questions grounded in the document, generate a quiz from its
# actual content, and save personal notes against it.

@app.post("/api/pdf-explainer/upload")
async def pdf_explainer_upload(file: UploadFile = File(...), child: str = Form("")):
    filename = file.filename or ""
    if Path(filename).suffix.lower() != ".pdf":
        raise HTTPException(status_code=400, detail="Only PDF files are supported")
    if not safety_filter.is_safe(filename):
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    contents = await file.read()
    try:
        record = pdf_explainer.upload(filename, contents, child=child)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    if record["summary"] and not safety_filter.is_safe(record["summary"]):
        pdf_explainer.delete_document(record["id"])
        raise HTTPException(status_code=400, detail="Upload rejected: unsafe content detected")

    return record


@app.get("/api/pdf-explainer")
def pdf_explainer_list(child: str = ""):
    return pdf_explainer.list_documents(child=child)


@app.get("/api/pdf-explainer/{doc_id}")
def pdf_explainer_get(doc_id: str):
    record = pdf_explainer.get_document(doc_id)
    if not record:
        raise HTTPException(status_code=404, detail="Document not found")
    return record


@app.delete("/api/pdf-explainer/{doc_id}")
def pdf_explainer_delete(doc_id: str):
    if not pdf_explainer.delete_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "deleted"}


@app.post("/api/pdf-explainer/{doc_id}/explain")
def pdf_explainer_explain(doc_id: str, body: dict):
    if not pdf_explainer.get_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    args = _tutor_level_args(body)
    explanation = pdf_explainer.explain(
        doc_id, level=args["level"], grade=args["grade"],
        age_group=args["age_group"], language=args["language"], difficulty=args["difficulty"],
    )
    return {"explanation": explanation}


@app.post("/api/pdf-explainer/{doc_id}/ask")
def pdf_explainer_ask(doc_id: str, body: dict):
    if not pdf_explainer.get_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    args = _tutor_level_args(body)
    question = safety_filter.sanitize(str(body.get("question", "")), strict=args["strict"])[:500]
    if not question:
        raise HTTPException(status_code=400, detail="question is required")
    answer = pdf_explainer.ask(doc_id, question, level=args["level"], grade=args["grade"])
    return {"answer": answer}


@app.post("/api/pdf-explainer/{doc_id}/quiz")
def pdf_explainer_quiz(doc_id: str, body: dict):
    if not pdf_explainer.get_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    args = _tutor_level_args(body)
    count = min(int(body.get("count", 5)), 10)
    questions = pdf_explainer.quiz(doc_id, count=count, level=args["level"], grade=args["grade"])
    return {"quiz": questions}


@app.get("/api/pdf-explainer/{doc_id}/notes")
def pdf_explainer_list_notes(doc_id: str, child: str = ""):
    if not pdf_explainer.get_document(doc_id):
        raise HTTPException(status_code=404, detail="Document not found")
    return pdf_explainer.list_notes(doc_id, child=child)


@app.post("/api/pdf-explainer/{doc_id}/notes")
def pdf_explainer_add_note(doc_id: str, body: dict):
    text = str(body.get("text", "")).strip()[:5000]
    if not text:
        raise HTTPException(status_code=400, detail="text is required")
    if not safety_filter.is_safe(text):
        raise HTTPException(status_code=400, detail="Note rejected: unsafe content detected")
    try:
        note = pdf_explainer.add_note(doc_id, text, child=str(body.get("child", "")))
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    return note


@app.delete("/api/pdf-explainer/notes/{note_id}")
def pdf_explainer_delete_note(note_id: str):
    if not pdf_explainer.delete_note(note_id):
        raise HTTPException(status_code=404, detail="Note not found")
    return {"status": "deleted"}


# ─── AI Lesson & Term Planner ───────────────────────────────────────────────
# Generate a sequential, term-length lesson plan for a subject and level,
# automatically scheduled across weekdays, and reschedule individual lessons
# afterwards. Used from the Professional Workspace.

@app.post("/api/lesson-planner/generate")
def lesson_planner_generate(body: dict):
    owner_id = str(body.get("owner_id", "")).strip()
    subject = str(body.get("subject", "")).strip()
    term_name = str(body.get("term_name", "")).strip()
    start_date = str(body.get("start_date", "")).strip()
    if not owner_id or not subject or not term_name or not start_date:
        raise HTTPException(status_code=400, detail="owner_id, subject, term_name and start_date are required")
    if not safety_filter.is_safe(f"{subject} {term_name}"):
        raise HTTPException(status_code=400, detail="Request rejected: unsafe content detected")

    args = _tutor_level_args(body)
    try:
        plan = lesson_planner.generate_plan(
            owner_id, subject, term_name, start_date,
            lesson_count=int(body.get("lesson_count", 10)),
            lessons_per_week=int(body.get("lessons_per_week", 3)),
            level=args["level"], grade=args["grade"],
            notes=safety_filter.sanitize(str(body.get("notes", "")), strict=args["strict"])[:600],
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return plan


@app.get("/api/lesson-planner")
def lesson_planner_list(owner_id: str = ""):
    return lesson_planner.list_plans(owner_id=owner_id)


@app.get("/api/lesson-planner/{plan_id}")
def lesson_planner_get(plan_id: str):
    plan = lesson_planner.get_plan(plan_id)
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    return plan


@app.delete("/api/lesson-planner/{plan_id}")
def lesson_planner_delete(plan_id: str):
    if not lesson_planner.delete_plan(plan_id):
        raise HTTPException(status_code=404, detail="Plan not found")
    return {"status": "deleted"}


@app.patch("/api/lesson-planner/{plan_id}/lessons/{lesson_id}")
def lesson_planner_reschedule(plan_id: str, lesson_id: str, body: dict):
    new_date = str(body.get("date", "")).strip()
    if not new_date:
        raise HTTPException(status_code=400, detail="date is required")
    try:
        plan = lesson_planner.reschedule_lesson(plan_id, lesson_id, new_date)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    if not plan:
        raise HTTPException(status_code=404, detail="Plan not found")
    return plan


class LocalLibraryScanRequest(BaseModel):
    folder: str
    analyse_books: bool = True
    max_files: int = 2000


@app.get("/api/course-providers")
def course_providers(query: str = ""):
    """Verified catalogue entry points plus subject-specific search links."""
    return course_catalog.catalogue(query)


@app.post("/api/local-library/scan")
def local_library_scan(request: LocalLibraryScanRequest):
    """Index owned local files in place; no source file is copied or uploaded."""
    try:
        return local_library.scan_folder(
            request.folder,
            analyse_books=request.analyse_books,
            max_files=request.max_files,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except PermissionError as exc:
        raise HTTPException(status_code=403, detail=f"Folder cannot be read: {exc}") from exc
    except OSError as exc:
        raise HTTPException(status_code=400, detail=f"Folder scan failed: {exc}") from exc


@app.post("/api/local-library/select-folder")
def local_library_select_folder():
    """Open the operating system folder picker on a locally installed desktop."""
    root = None
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        root.update()
        folder = filedialog.askdirectory(
            parent=root,
            title="Select a learning resources folder",
            mustexist=True,
        )
        return {"folder": folder or "", "cancelled": not bool(folder)}
    except Exception as exc:
        raise HTTPException(
            status_code=503,
            detail="The system folder picker is unavailable; paste the folder path instead",
        ) from exc
    finally:
        if root is not None:
            try:
                root.destroy()
            except Exception:
                pass


@app.get("/api/local-library")
def local_library_list(category: str = "", query: str = "", limit: int = 500):
    if category and category not in local_library.CATEGORY_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Unknown local-library category")
    return {
        "categories": sorted(local_library.CATEGORY_EXTENSIONS),
        "files": local_library.list_files(category=category, query=query, limit=limit),
    }


@app.get("/api/local-library/files/{file_id}")
def local_library_open(file_id: str):
    result = local_library.get_file(file_id)
    if not result:
        raise HTTPException(status_code=404, detail="Local file is unavailable or has moved")
    record, path = result
    return FileResponse(
        path,
        filename=record["filename"],
        content_disposition_type="inline",
    )


class PaintingSaveRequest(BaseModel):
    title: str = "Untitled"
    image: str  # data URL or raw base64 PNG
    id: str | None = None


@app.get("/api/paintings/{child}")
def list_paintings(child: str):
    _require_child(child)
    return {"paintings": paintings_store.list_paintings(child)}


@app.post("/api/paintings/{child}")
def save_painting(child: str, body: PaintingSaveRequest):
    _require_child(child)
    try:
        record = paintings_store.save_painting(child, body.title, body.image, body.id)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    return record


@app.get("/api/paintings/{child}/{painting_id}/image")
def get_painting_image(child: str, painting_id: str):
    _require_child(child)
    path = paintings_store.get_painting_path(child, painting_id)
    if not path or not path.exists():
        raise HTTPException(status_code=404, detail="Painting not found")
    with open(path, "rb") as f:
        data = f.read()
    return StreamingResponse(BytesIO(data), media_type="image/png")


@app.delete("/api/paintings/{child}/{painting_id}")
def delete_painting(child: str, painting_id: str):
    _require_child(child)
    if not paintings_store.delete_painting(child, painting_id):
        raise HTTPException(status_code=404, detail="Painting not found")
    return {"status": "deleted"}


@app.get("/api/safe-music")
def safe_music():
    with open(SAFE_DIR / "safe_songs.json", encoding="utf-8") as f:
        data = json.load(f)
    songs = [s for s in data["songs"] if s.get("safe")]
    return songs


@app.get("/api/activity-log/{child}")
def activity_log(child: str):
    _require_child(child)
    return get_activity_log(child)


@app.get("/api/safe-channels")
def safe_channels():
    with open(SAFE_DIR / "safe_channels.json", encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/sing-along-songs")
def sing_along_songs():
    with open(SAFE_DIR / "sing_along_songs.json", encoding="utf-8") as f:
        data = json.load(f)
    return [s for s in data["songs"] if s.get("safe")]


@app.get("/api/profiles")
def profiles():
    return get_all_profiles()


# ─── User Management (Parent-only) ────────────────────────────────────────────

@app.get("/api/users")
def list_users():
    from app.storage import _load_users
    data = _load_users()
    users = [{"name": n, "role": "child"} for n in data["children"]] + \
            [{"name": n, "role": "parent"} for n in data["parents"]]
    return {"users": users}


class UserCreate(BaseModel):
    name: str
    role: str  # "child" or "parent"

class UserRename(BaseModel):
    new_name: str

@app.post("/api/users")
def create_user(body: UserCreate):
    raise HTTPException(status_code=403, detail="EduAI_Pro is configured for the single administrator Shovan")


@app.put("/api/users/{name}")
def update_user(name: str, body: UserRename):
    new_name = body.new_name.strip()
    if not new_name or not new_name.replace(" ", "").isalnum():
        raise HTTPException(status_code=400, detail="Name must be non-empty and alphanumeric")
    try:
        data = rename_user(name, new_name)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"ok": True, "users": data}


@app.delete("/api/users/{name}")
def remove_user(name: str):
    try:
        data = delete_user(name)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return {"ok": True, "users": data}


@app.get("/api/web-search")
def web_search_endpoint(q: str):
    query = q.strip()
    if not query:
        return []
    try:
        return web_search(query)
    except SearchNotConfigured as exc:
        raise HTTPException(status_code=501, detail=str(exc)) from exc


@app.post("/api/curate-resource")
def curate_resource_endpoint(payload: dict):
    try:
        standard = int(payload["standard"])
        subject = payload["subject"]
        resource_type = payload["resource_type"]
        resource = payload["resource"]
    except (KeyError, TypeError, ValueError) as exc:
        raise HTTPException(status_code=422, detail="standard, subject, resource_type, resource are required") from exc

    if resource_type not in CURATE_RESOURCE_KEYS:
        raise HTTPException(status_code=422, detail=f"resource_type must be one of {CURATE_RESOURCE_KEYS}")

    try:
        saved = curate_resource(standard, subject, resource_type, resource)
    except CurationError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return saved


RESOURCE_KEYS = (
    "books",
    "video_resources",
    "text_resources",
    "cartoon_videos",
    "infographics",
    "textbooks",
    "audio_resources",
    "comics",
    "drawing_activities",
    "info_cards",
    "news_resources",
)


@app.post("/api/resource/download")
def download_resource(payload: dict):
    """Download a single resource's saved text (and any media link) as a file.

    Works for any resource that has a 'body' (full saved article/text) and/or
    'image' field already stored locally — it never fetches anything live.
    """
    title = str(payload.get("title") or "Untitled resource")
    body = str(payload.get("body") or payload.get("description") or "").strip()
    url = str(payload.get("url") or "")
    source = str(payload.get("source") or "")
    image = str(payload.get("image") or "")
    fmt = str(payload.get("format") or "txt").lower()

    if not safety_filter.is_safe(title) or (body and not safety_filter.is_safe(body)):
        raise HTTPException(status_code=400, detail="Download rejected: unsafe content detected")

    safe_name = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)[:80].strip() or "resource"

    if fmt == "docx":
        document = Document()
        document.add_heading(title, level=1)
        if source:
            document.add_paragraph(f"Source: {source}")
        if url:
            document.add_paragraph(f"Original link: {url}")
        if body:
            document.add_paragraph(body)
        else:
            document.add_paragraph(
                "No saved article text is available for this resource yet — only its source link."
            )
        if image:
            document.add_paragraph(f"Media: {image}")
        buffer = BytesIO()
        document.save(buffer)
        return StreamingResponse(
            BytesIO(buffer.getvalue()),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_name}.docx"'},
        )

    lines = [title, "=" * len(title), ""]
    if source:
        lines.append(f"Source: {source}")
    if url:
        lines.append(f"Original link: {url}")
    lines.append("")
    lines.append(body or "No saved article text is available for this resource yet — only its source link.")
    if image:
        lines.append("")
        lines.append(f"Media: {image}")
    text = "\n".join(lines)

    return StreamingResponse(
        BytesIO(text.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.txt"'},
    )


@app.get("/api/search/{standard}")
def search_grade(standard: int, q: str):
    path = SYLLABUS_DIR / f"grade{standard}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grade {standard} not available yet")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)

    query = q.strip().lower()
    if not query:
        return []

    results = []
    for subject_name, subject in data.get("subjects", {}).items():
        for key in RESOURCE_KEYS:
            for resource in subject.get(key, []):
                if resource.get("safe") is not True:
                    continue
                haystack = " ".join(
                    str(v) for v in (resource.get("title"), resource.get("description"))
                    if v
                ).lower()
                if query in haystack:
                    results.append(
                        _sanitize_json({**resource, "subject": subject_name, "resource_type": key})
                    )
    return results


# ─── AI Tutor ────────────────────────────────────────────────────────────────
# The tutor serves school, college, undergraduate, master's, and adult
# self-learners. Callers may pass either the legacy numeric `grade`, or the
# new `level` code (e.g. "UG2", "M1"); `age_group`, `language`, and
# `difficulty` further tailor the response. Safety filtering relaxes to
# hard-blocks-only (never fully off) once the level is college/UG/master's.

def _tutor_level_args(body: dict) -> dict:
    level = body.get("level")
    grade = int(body.get("grade", 1))
    strict = not (level and levels_module.is_adult_level(level))
    return {
        "level": levels_module.normalize_level_id(level) if level else None,
        "grade": grade,
        "strict": strict,
        "age_group": str(body.get("age_group", "")),
        "language": str(body.get("language", "")),
        "difficulty": str(body.get("difficulty", "")),
    }


@app.post("/api/ai-tutor/ask")
def tutor_ask(body: dict):
    args = _tutor_level_args(body)
    question = safety_filter.sanitize(str(body.get("question", "")), strict=args["strict"])[:500]
    subject = str(body.get("subject", ""))
    context = str(body.get("context", ""))[:600]
    if not question:
        raise HTTPException(status_code=400, detail="question is required")
    answer = ai_tutor.ask(
        question, grade=args["grade"], subject=subject, context=context, level=args["level"],
        age_group=args["age_group"], language=args["language"], difficulty=args["difficulty"],
    )
    return {"answer": answer}


@app.post("/api/ai-tutor/grounded")
def tutor_grounded(body: dict):
    question = str(body.get("question", "")).strip()[:4000]
    user_id = str(body.get("user_id", "")).strip()
    level = levels_module.normalize_level_id(body.get("level", "1"))
    if not question or not user_id:
        raise HTTPException(status_code=400, detail="question and user_id are required")
    if not levels_module.is_valid_level(level):
        raise HTTPException(status_code=400, detail="Unknown academic level")
    try:
        return ai_reliability.grounded_answer(
            user_id=user_id,
            question=question,
            level_id=level,
            subject=str(body.get("subject", "")),
            difficulty=str(body.get("difficulty", "")),
            mode=str(body.get("mode", "direct")),
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc


@app.post("/api/ai-tutor/explain")
def tutor_explain(body: dict):
    args = _tutor_level_args(body)
    concept = safety_filter.sanitize(str(body.get("concept", "")), strict=args["strict"])[:300]
    subject = str(body.get("subject", ""))
    if not concept:
        raise HTTPException(status_code=400, detail="concept is required")
    explanation = ai_tutor.explain_concept(
        concept, grade=args["grade"], subject=subject, level=args["level"],
        age_group=args["age_group"], language=args["language"], difficulty=args["difficulty"],
    )
    return {"explanation": explanation}


@app.post("/api/ai-tutor/flashcards")
def tutor_flashcards(body: dict):
    args = _tutor_level_args(body)
    topic = safety_filter.sanitize(str(body.get("topic", "")), strict=args["strict"])[:200]
    subject = str(body.get("subject", ""))
    count = min(int(body.get("count", 8)), 20)
    if not topic:
        raise HTTPException(status_code=400, detail="topic is required")
    cards = ai_tutor.generate_flashcards(topic, grade=args["grade"], subject=subject, count=count, level=args["level"])
    return {"flashcards": cards}


@app.post("/api/ai-tutor/quiz")
def tutor_quiz(body: dict):
    args = _tutor_level_args(body)
    topic = safety_filter.sanitize(str(body.get("topic", "")), strict=args["strict"])[:200]
    subject = str(body.get("subject", ""))
    count = min(int(body.get("count", 5)), 10)
    if not topic:
        raise HTTPException(status_code=400, detail="topic is required")
    questions = ai_tutor.generate_quiz(topic, grade=args["grade"], subject=subject, count=count, level=args["level"])
    return {"quiz": questions}


@app.post("/api/ai-tutor/study-plan")
def tutor_study_plan(body: dict):
    args = _tutor_level_args(body)
    subject = safety_filter.sanitize(str(body.get("subject", "")), strict=args["strict"])[:100]
    days = min(int(body.get("days", 7)), 30)
    if not subject:
        raise HTTPException(status_code=400, detail="subject is required")
    plan = ai_tutor.make_study_plan(subject, grade=args["grade"], days=days, level=args["level"])
    return {"plan": plan}


# ─── Language Academy ────────────────────────────────────────────────────────

LANG_DIR = BASE_DIR / "data" / "language_academy"


@app.get("/api/languages")
def list_languages():
    path = LANG_DIR / "languages.json"
    if not path.exists():
        return {"languages": []}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/languages/{code}")
def get_language(code: str):
    path = LANG_DIR / f"vocab_{code}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Language '{code}' not available")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if "categories" in data and "vocabulary" not in data:
        data["vocabulary"] = data["categories"]
    lang_path = LANG_DIR / "languages.json"
    if lang_path.exists():
        with open(lang_path, encoding="utf-8") as lf:
            langs = json.load(lf).get("languages", [])
        meta = next((l for l in langs if l.get("code") == code), {})
        for k, v in meta.items():
            data.setdefault(k, v)
    return data


@app.get("/api/languages/{code}/quiz")
def get_language_quiz(code: str):
    import random
    path = LANG_DIR / f"vocab_{code}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Language '{code}' not available")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    vocab = data.get("vocabulary", [])
    if not vocab:
        return {"quiz": [], "language": data.get("language", code)}
    sample = random.sample(vocab, min(10, len(vocab)))
    quiz = []
    for item in sample:
        wrong_pool = [v for v in vocab if v["word"] != item["word"]]
        distractors = [w["translation"] for w in random.sample(wrong_pool, min(3, len(wrong_pool)))]
        options = distractors + [item["translation"]]
        random.shuffle(options)
        quiz.append({
            "question": f"What does '{item['word']}' mean?",
            "options": options,
            "answer": options.index(item["translation"]),
        })
    return {"quiz": quiz, "language": data.get("language", code)}


@app.get("/api/languages/{code}/sentences")
def get_language_sentences(code: str):
    path = LANG_DIR / f"sentences_{code}.json"
    if not path.exists():
        return {"sentences": [], "language": code}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ─── Assessment Centre ───────────────────────────────────────────────────────

ASSESSMENT_DIR = BASE_DIR / "data" / "assessment"


@app.get("/api/assessment/age-groups")
def list_age_groups():
    path = ASSESSMENT_DIR / "assessments.json"
    if not path.exists():
        return {"age_groups": []}
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return {
        "age_groups": [
            {"id": k, "label": v["label"], "description": v["description"]}
            for k, v in data.get("age_groups", {}).items()
        ],
        "disclaimer": data.get("disclaimer", "")
    }


@app.get("/api/assessment/{age_group}")
def get_assessment(age_group: str):
    path = ASSESSMENT_DIR / "assessments.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Assessment data not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    group = data.get("age_groups", {}).get(age_group)
    if not group:
        raise HTTPException(status_code=404, detail=f"Age group '{age_group}' not found")
    return {**group, "disclaimer": data.get("disclaimer", "")}


@app.post("/api/assessment/{child}/submit")
def submit_assessment(child: str, body: dict):
    _require_child(child)
    age_group = body.get("age_group", "")
    answers = body.get("answers", {})  # {"{section_index}-{question_index}": chosen_option_index}

    path = ASSESSMENT_DIR / "assessments.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Assessment data not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    group = data.get("age_groups", {}).get(age_group)
    if not group:
        raise HTTPException(status_code=404, detail=f"Age group '{age_group}' not found")
    skill_map = data.get("skill_recommendations", {})

    # Grade server-side from the assessment's own answer key rather than trusting
    # a client-submitted score, so this stays correct regardless of what the
    # frontend sends and can't be spoofed.
    skill_correct: dict[str, int] = {}
    skill_total: dict[str, int] = {}
    score = 0
    total = 0
    for si, section in enumerate(group.get("sections", [])):
        for qi, q in enumerate(section.get("questions", [])):
            total += 1
            skill = q.get("skill", "general")
            skill_total[skill] = skill_total.get(skill, 0) + 1
            if answers.get(f"{si}-{qi}") == q.get("answer"):
                score += 1
                skill_correct[skill] = skill_correct.get(skill, 0) + 1

    strengths = []
    areas_to_develop = []
    weak_skills = []  # raw skill keys (not humanized) so a retake can filter questions by them
    recommendations: list[str] = []
    seen_subjects = set()
    for skill, s_total in skill_total.items():
        label = skill.replace("_", " ")
        if skill_correct.get(skill, 0) == s_total:
            strengths.append(label)
            for subj in skill_map.get(skill, []):
                if subj not in seen_subjects:
                    recommendations.append(subj)
                    seen_subjects.add(subj)
        else:
            areas_to_develop.append(label)
            weak_skills.append(skill)

    percentage = round(score / total * 100) if total else 0
    badge = None
    if total > 0 and score / total >= 0.8:
        badge = f"assessment-{age_group}-distinction"
        save_progress(child, {"badges": [badge]})

    append_activity(child, {
        "type": "assessment",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "age_group": age_group,
        "score": score,
        "total": total,
        "percentage": percentage,
        "badge": badge,
        "strengths": strengths,
        "areas_to_develop": areas_to_develop,
        "weak_skills": weak_skills,
        "recommendations": recommendations[:8],
    })
    return {
        "score": score,
        "total": total,
        "percentage": percentage,
        "badge": badge,
        "strengths": strengths,
        "areas_to_develop": areas_to_develop,
        "weak_skills": weak_skills,
        "recommendations": recommendations[:8],
        "message": "Well done! Keep learning and growing." if total and score / total >= 0.6 else "Great effort! Review the topics you found tricky and try again."
    }


@app.get("/api/assessment/{child}/history")
def assessment_history(child: str):
    """Every past assessment attempt for this child, newest first -- lets the
    Assessment Centre show a learning-profile history over time instead of
    only the most recent result."""
    _require_child(child)
    attempts = [a for a in get_activity_log(child) if a.get("type") == "assessment"]
    attempts.sort(key=lambda a: a.get("timestamp", ""), reverse=True)
    return {"attempts": attempts}


@app.get("/api/assessment/{age_group}/retake")
def assessment_retake(age_group: str, weak_skills: str = ""):
    """A retake assessment covering only the given comma-separated skill keys
    (as returned by a previous submit's `weak_skills`), so a child can focus
    practice on what they actually got wrong last time instead of retaking
    the whole assessment. Falls back to the full assessment if no skills (or
    only unrecognised skills) are given."""
    path = ASSESSMENT_DIR / "assessments.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Assessment data not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    group = data.get("age_groups", {}).get(age_group)
    if not group:
        raise HTTPException(status_code=404, detail=f"Age group '{age_group}' not found")

    requested = {s.strip() for s in weak_skills.split(",") if s.strip()}
    if not requested:
        return {**group, "disclaimer": data.get("disclaimer", ""), "is_retake": False}

    filtered_sections = []
    for section in group.get("sections", []):
        questions = [q for q in section.get("questions", []) if q.get("skill") in requested]
        if questions:
            filtered_sections.append({**section, "questions": questions})

    if not filtered_sections:
        # None of the requested skills matched this age group's questions --
        # fall back to the full assessment rather than returning an empty one.
        return {**group, "disclaimer": data.get("disclaimer", ""), "is_retake": False}

    return {
        **group,
        "sections": filtered_sections,
        "disclaimer": data.get("disclaimer", ""),
        "is_retake": True,
        "label": f"{group.get('label', '')} — Retake: Focus Areas",
    }


# ─── Grammar Academy ─────────────────────────────────────────────────────────

GRAMMAR_DIR = BASE_DIR / "data" / "grammar"


@app.get("/api/grammar")
def get_grammar_curriculum():
    path = GRAMMAR_DIR / "grammar_curriculum.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Grammar curriculum not found")
    with open(path, encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/grammar/language/{lang_code}/{level}")
def get_grammar_language_level(lang_code: str, level: str):
    path = GRAMMAR_DIR / f"grammar_{lang_code}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grammar for '{lang_code}' not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    level_data = data.get("levels", {}).get(level)
    if not level_data:
        raise HTTPException(status_code=404, detail=f"Level '{level}' not found")
    return level_data


@app.get("/api/grammar/language/{lang_code}")
def get_grammar_language(lang_code: str):
    path = GRAMMAR_DIR / f"grammar_{lang_code}.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail=f"Grammar for '{lang_code}' not found")
    with open(path, encoding="utf-8") as f:
        return json.load(f)


@app.get("/api/grammar/{level}")
def get_grammar_level(level: str):
    path = GRAMMAR_DIR / "grammar_curriculum.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Grammar curriculum not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    level_data = data.get("levels", {}).get(level)
    if not level_data:
        raise HTTPException(status_code=404, detail=f"Level '{level}' not found")
    return level_data


# ─── Countries ───────────────────────────────────────────────────────────────

COUNTRIES_DIR = BASE_DIR / "data" / "countries"


@app.get("/api/countries")
def list_countries():
    path = COUNTRIES_DIR / "countries.json"
    if not path.exists():
        return {"countries": [], "total": 0}
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return {"countries": data.get("countries", []), "total": len(data.get("countries", []))}


@app.get("/api/countries/{code}")
def get_country(code: str):
    path = COUNTRIES_DIR / "countries.json"
    if not path.exists():
        raise HTTPException(status_code=404, detail="Countries data not found")
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    country = next((c for c in data.get("countries", []) if c.get("code", "").upper() == code.upper()), None)
    if not country:
        raise HTTPException(status_code=404, detail=f"Country '{code}' not found")
    return country


# ─── Parent Dashboard ─────────────────────────────────────────────────────────

import uuid as _uuid

@app.get("/api/parent/homework/{child}")
def get_child_homework(child: str):
    _require_child(child)
    return {"homework": get_homework(child)}


@app.post("/api/parent/homework/{child}")
def add_homework(child: str, body: dict):
    _require_child(child)
    items = get_homework(child)
    item = {
        "id": str(_uuid.uuid4())[:8],
        "subject": body.get("subject", ""),
        "title": body.get("title", ""),
        "due_date": body.get("due_date", ""),
        "status": body.get("status", "pending"),
        "notes": body.get("notes", ""),
    }
    items.append(item)
    save_homework(child, items)
    return item


@app.patch("/api/parent/homework/{child}/{hw_id}")
def update_homework(child: str, hw_id: str, body: dict):
    _require_child(child)
    items = get_homework(child)
    for item in items:
        if item["id"] == hw_id:
            item.update({k: v for k, v in body.items() if k != "id"})
            save_homework(child, items)
            return item
    raise HTTPException(status_code=404, detail="Homework item not found")


@app.delete("/api/parent/homework/{child}/{hw_id}")
def delete_homework(child: str, hw_id: str):
    _require_child(child)
    items = [i for i in get_homework(child) if i["id"] != hw_id]
    save_homework(child, items)
    return {"status": "deleted"}


@app.get("/api/parent/reading-log/{child}")
def get_child_reading_log(child: str):
    _require_child(child)
    log = get_reading_log(child)
    total_pages = sum(e.get("pages", 0) for e in log)
    total_minutes = sum(e.get("duration_mins", 0) for e in log)
    return {"log": log, "total_pages": total_pages, "total_minutes": total_minutes}


@app.post("/api/parent/reading-log/{child}")
def add_reading_entry(child: str, body: dict):
    _require_child(child)
    from datetime import date as _date
    entry = {
        "date": body.get("date", _date.today().isoformat()),
        "book": body.get("book", ""),
        "author": body.get("author", ""),
        "pages": int(body.get("pages", 0)),
        "duration_mins": int(body.get("duration_mins", 0)),
        "notes": body.get("notes", ""),
    }
    log = append_reading_entry(child, entry)
    return {"entry": entry, "total_entries": len(log)}


@app.get("/api/parent/screen-time/{child}")
def get_child_screen_time(child: str):
    _require_child(child)
    data = get_screen_time(child)
    total = sum(data.values())
    sorted_days = sorted(data.items(), reverse=True)[:14]
    return {"daily": dict(sorted_days), "total_minutes": total}


@app.post("/api/parent/screen-time/{child}/add")
def record_screen_time(child: str, body: dict):
    _require_child(child)
    minutes = int(body.get("minutes", 0))
    date_str = body.get("date", None)
    updated = add_screen_time(child, minutes, date_str)
    return {"updated": updated}


@app.get("/api/parent/weekly-report/{child}")
def get_weekly_report(child: str):
    _require_child(child)
    from datetime import date as _date, timedelta as _td
    progress = get_progress(child)
    reading = get_reading_log(child)
    screen = get_screen_time(child)
    homework = get_homework(child)
    activity = get_activity_log(child)

    today = _date.today()
    week_ago = (today - _td(days=7)).isoformat()

    recent_reading = [e for e in reading if e.get("date", "") >= week_ago]
    recent_screen = {k: v for k, v in screen.items() if k >= week_ago}
    recent_activity = [a for a in activity if a.get("timestamp", "") >= week_ago]
    pending_hw = [h for h in homework if h.get("status") == "pending"]
    done_hw = [h for h in homework if h.get("status") == "done"]

    return {
        "child": child,
        "week_of": week_ago,
        "lesson_streak": progress.get("lesson_streak", 0),
        "badges_earned": progress.get("badges", []),
        "subjects_studied": list(progress.get("completed_lessons", {}).keys()),
        "reading_sessions": len(recent_reading),
        "reading_pages": sum(e.get("pages", 0) for e in recent_reading),
        "reading_minutes": sum(e.get("duration_mins", 0) for e in recent_reading),
        "screen_time_minutes": sum(recent_screen.values()),
        "activities_completed": len(recent_activity),
        "homework_pending": len(pending_hw),
        "homework_done": len(done_hw),
        "scores": progress.get("scores", {}),
    }


# ── English Vocabulary Academy ──────────────────────────────────────────────
_VOCAB_PATH = Path(__file__).parent.parent / "data" / "vocabulary" / "vocab_academy.json"

def _load_vocab() -> dict:
    with open(_VOCAB_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/vocabulary")
def vocabulary_overview():
    data = _load_vocab()
    summary = []
    for level_key, level in data["levels"].items():
        cats = level["categories"]
        total_words = sum(len(v) if isinstance(v, list) else len(v.get("words", [])) for v in cats.values())
        summary.append({
            "id": level_key,
            "label": level["label"],
            "word_count": level.get("word_count", total_words),
            "categories": list(level["categories"].keys()),
        })
    return {"title": data["title"], "description": data["description"], "levels": summary}

@app.get("/api/vocabulary/search")
def vocabulary_search(q: str = ""):
    if not q or len(q) < 2:
        return {"results": []}
    data = _load_vocab()
    q_lower = q.lower()
    results = []
    for level_key, level in data["levels"].items():
        for cat_name, cat in level["categories"].items():
            for word_entry in (cat if isinstance(cat, list) else cat.get("words", [])):
                w = word_entry.get("word", "")
                m = word_entry.get("meaning", "")
                if q_lower in w.lower() or q_lower in m.lower():
                    results.append({
                        "word": w,
                        "meaning": m,
                        "example": word_entry.get("example", ""),
                        "synonyms": word_entry.get("synonyms", []),
                        "antonyms": word_entry.get("antonyms", []),
                        "level": level_key,
                        "category": cat_name,
                    })
    return {"results": results[:50]}

@app.get("/api/vocabulary/{level}")
def vocabulary_level(level: str):
    data = _load_vocab()
    if level not in data["levels"]:
        raise HTTPException(status_code=404, detail="Level not found")
    lv = data["levels"][level]
    return {"id": level, **lv}

@app.get("/api/vocabulary/{level}/quiz")
def vocabulary_quiz(level: str):
    data = _load_vocab()
    if level not in data["levels"]:
        raise HTTPException(status_code=404, detail="Level not found")
    return {"level": level, "quiz": data["levels"][level].get("quiz", [])}



# ── STEM Laboratory ──────────────────────────────────────────────────────────
_STEM_PATH = Path(__file__).parent.parent / "data" / "stem_lab" / "stem_lab.json"

def _load_stem() -> dict:
    with open(_STEM_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/stem-lab")
def stem_lab_overview():
    data = _load_stem()
    disciplines = []
    for key, disc in data["disciplines"].items():
        disciplines.append({
            "id": key,
            "label": disc["label"],
            "emoji": disc["emoji"],
            "colour": disc["colour"],
            "description": disc["description"],
            "experiment_count": len(disc.get("experiments", [])),
        })
    return {"title": data["title"], "description": data["description"], "disciplines": disciplines}

@app.get("/api/stem-lab/{discipline}")
def stem_lab_discipline(discipline: str):
    data = _load_stem()
    if discipline not in data["disciplines"]:
        raise HTTPException(status_code=404, detail="Discipline not found")
    return {"id": discipline, **data["disciplines"][discipline]}

@app.get("/api/stem-lab/{discipline}/{experiment_id}")
def stem_lab_experiment(discipline: str, experiment_id: str):
    data = _load_stem()
    if discipline not in data["disciplines"]:
        raise HTTPException(status_code=404, detail="Discipline not found")
    exps = data["disciplines"][discipline].get("experiments", [])
    for exp in exps:
        if exp["id"] == experiment_id:
            return exp
    raise HTTPException(status_code=404, detail="Experiment not found")


# ── Non-Fiction Library ──────────────────────────────────────────────────────
_NONFICTION_PATH = Path(__file__).parent.parent / "data" / "nonfiction_library" / "nonfiction.json"

def _load_nonfiction() -> dict:
    with open(_NONFICTION_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/nonfiction")
def nonfiction_overview():
    data = _load_nonfiction()
    cats = []
    for key, cat in data["categories"].items():
        cats.append({
            "id": key,
            "label": cat["label"],
            "emoji": cat["emoji"],
            "book_count": len(cat.get("books", [])),
        })
    return {"title": data["title"], "description": data["description"], "categories": cats}

@app.get("/api/nonfiction/{category}")
def nonfiction_category(category: str):
    data = _load_nonfiction()
    if category not in data["categories"]:
        raise HTTPException(status_code=404, detail="Category not found")
    return {"id": category, **data["categories"][category]}

@app.get("/api/nonfiction/{category}/{book_id}")
def nonfiction_book(category: str, book_id: str):
    data = _load_nonfiction()
    if category not in data["categories"]:
        raise HTTPException(status_code=404, detail="Category not found")
    for book in data["categories"][category].get("books", []):
        if book["id"] == book_id:
            return book
    raise HTTPException(status_code=404, detail="Book not found")


# ── Practical Skills Academy ─────────────────────────────────────────────────
_PRACTICAL_PATH = Path(__file__).parent.parent / "data" / "practical_skills" / "practical_skills.json"

def _load_practical() -> dict:
    with open(_PRACTICAL_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/practical-skills")
def practical_skills_overview():
    data = _load_practical()
    pathways = []
    for key, pw in data["pathways"].items():
        pathways.append({
            "id": key,
            "label": pw.get("label", key.replace("_", " ").title()),
            "emoji": pw.get("emoji", "📚"),
            "level_count": len(pw.get("levels", {})),
            "module_count": len(pw.get("modules", [])),
        })
    return {"title": data["title"], "description": data["description"], "pathways": pathways}

@app.get("/api/practical-skills/{pathway}")
def practical_skills_pathway(pathway: str):
    data = _load_practical()
    if pathway not in data["pathways"]:
        raise HTTPException(status_code=404, detail="Pathway not found")
    return {"id": pathway, **data["pathways"][pathway]}

@app.get("/api/practical-skills/{pathway}/{level}")
def practical_skills_level(pathway: str, level: str):
    data = _load_practical()
    if pathway not in data["pathways"]:
        raise HTTPException(status_code=404, detail="Pathway not found")
    levels = data["pathways"][pathway].get("levels", {})
    if level not in levels:
        raise HTTPException(status_code=404, detail="Level not found")
    return {"pathway": pathway, "level": level, **levels[level]}


# ── Virtual Museum ────────────────────────────────────────────────────────────
_MUSEUM_PATH = Path(__file__).parent.parent / "data" / "virtual_museum" / "museum.json"
_MUSEUM_OBJECTS_PATH = Path(__file__).parent.parent / "data" / "museum_objects.json"
_MUSEUM_IMAGE_CACHE = Path(__file__).parent.parent / "data" / "museum_resource" / "images"
_MUSEUM_IMAGE_CACHE.mkdir(parents=True, exist_ok=True)
_MUSEUM_THUMB_CACHE_DIR = Path(__file__).parent.parent / "data" / "museum_thumbnail_cache"
_MUSEUM_THUMB_CACHE_DIR.mkdir(parents=True, exist_ok=True)
_MUSEUM_THUMB_TTL_SECONDS = 5 * 24 * 3600  # Wikipedia thumbnails rarely change; a few days is plenty

def _load_museum() -> dict:
    with open(_MUSEUM_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/museum")
def museum_overview():
    data = _load_museum()
    galleries = []
    total_objects = 0
    for key, gallery in data["galleries"].items():
        object_count = len(gallery.get("objects", []))
        total_objects += object_count
        galleries.append({
            "id": key,
            "label": gallery["label"],
            "emoji": gallery["emoji"],
            "object_count": object_count,
        })
    return {
        "title": data["title"],
        "description": data["description"],
        "total_objects": total_objects,
        "galleries": galleries,
    }

@app.get("/api/museum/search")
def museum_search(q: str = ""):
    if not q or len(q) < 2:
        return {"results": []}
    data = _load_museum()
    q_lower = q.lower()
    results = []
    for gallery_id, gallery in data["galleries"].items():
        for obj in gallery.get("objects", []):
            if (q_lower in obj.get("name", "").lower() or
                q_lower in obj.get("origin", "").lower() or
                q_lower in obj.get("description", "").lower() or
                any(q_lower in s.lower() for s in obj.get("related_subjects", []))):
                results.append({**obj, "gallery": gallery_id, "gallery_label": gallery["label"]})
    return {"results": results[:20]}


def _museum_thumb_cache_path(wiki_title: str) -> Path:
    # Hash the title so arbitrary query input can't escape the cache directory
    # or collide with filesystem-unsafe characters.
    digest = hashlib.sha1(wiki_title.encode("utf-8")).hexdigest()
    return _MUSEUM_THUMB_CACHE_DIR / f"{digest}.json"


@app.get("/api/museum/thumbnail")
def museum_thumbnail(wiki_title: str = ""):
    """Server-side cache/proxy for Wikipedia page-summary thumbnails.

    Museum object thumbnails come from Wikipedia's REST summary API. Fetching
    that live from every browser, on every view, means the same handful of
    popular objects get re-fetched from Wikipedia over and over across
    sessions and users. This caches the resolved thumbnail URL to a small
    JSON record on disk (keyed by a hash of wiki_title) with a TTL, so repeat
    requests for the same object are served locally instead of hitting
    Wikipedia again.
    """
    wiki_title = (wiki_title or "").strip()
    if not wiki_title:
        raise HTTPException(status_code=400, detail="wiki_title is required")

    cache_path = _museum_thumb_cache_path(wiki_title)
    now = time.time()
    cached = None
    if cache_path.exists():
        try:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
        except Exception:
            cached = None
        if cached and now - cached.get("fetched_at", 0) < _MUSEUM_THUMB_TTL_SECONDS:
            return {"wiki_title": wiki_title, "thumbnail_url": cached.get("thumbnail_url"), "cached": True}

    thumbnail_url = None
    fetch_ok = False
    encoded = quote(wiki_title.replace(" ", "_"))
    url = f"https://en.wikipedia.org/api/rest_v1/page/summary/{encoded}"
    try:
        req = Request(url, headers={
            "Api-User-Agent": "EduAI/1.0 (educational; contact@eduai.app)",
            "User-Agent": "EduAI/1.0 (educational; contact@eduai.app)",
        })
        with urlopen(req, timeout=5) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        thumbnail_url = (payload.get("thumbnail") or {}).get("source") or (payload.get("originalimage") or {}).get("source")
        fetch_ok = True
    except Exception:
        fetch_ok = False

    if not fetch_ok and cached:
        # Wikipedia is unreachable/rate-limited right now -- serve the stale
        # cached value rather than nothing, and don't overwrite it.
        return {"wiki_title": wiki_title, "thumbnail_url": cached.get("thumbnail_url"), "cached": True, "stale": True}

    record = {"wiki_title": wiki_title, "thumbnail_url": thumbnail_url, "fetched_at": now}
    try:
        cache_path.write_text(json.dumps(record), encoding="utf-8")
    except Exception:
        pass
    return {"wiki_title": wiki_title, "thumbnail_url": thumbnail_url, "cached": False}


@app.get("/api/museum/{gallery}")
def museum_gallery(gallery: str):
    data = _load_museum()
    if gallery not in data["galleries"]:
        raise HTTPException(status_code=404, detail="Gallery not found")
    return {"id": gallery, **data["galleries"][gallery]}

@app.get("/api/museum/{gallery}/{object_id}")
def museum_object(gallery: str, object_id: str):
    data = _load_museum()
    if gallery not in data["galleries"]:
        raise HTTPException(status_code=404, detail="Gallery not found")
    for obj in data["galleries"][gallery].get("objects", []):
        if obj["id"] == object_id:
            return obj
    raise HTTPException(status_code=404, detail="Object not found")


# ── CMA Open-Access Museum Objects ──────────────────────────────────────────

@lru_cache(maxsize=1)
def _load_museum_objects() -> list:
    if not _MUSEUM_OBJECTS_PATH.exists():
        return []
    with open(_MUSEUM_OBJECTS_PATH, encoding="utf-8") as f:
        d = json.load(f)
    return d.get("objects", d) if isinstance(d, dict) else d


@app.get("/api/museum-objects")
def museum_objects_list(
    q: str = "",
    type_filter: str = "",
    culture: str = "",
    page: int = 1,
    per_page: int = 40,
):
    objs = _load_museum_objects()
    if q:
        ql = q.lower()
        objs = [o for o in objs if ql in (o.get("title") or "").lower()
                or ql in (o.get("creator") or "").lower()
                or ql in (o.get("description") or "").lower()]
    if type_filter:
        objs = [o for o in objs if (o.get("type") or "").lower() == type_filter.lower()]
    if culture:
        objs = [o for o in objs if any(culture.lower() in (c.lower() if isinstance(c, str) else "") for c in (o.get("culture") or []))]
    total = len(objs)
    start = (page - 1) * per_page
    slice_ = objs[start: start + per_page]
    return {"total": total, "page": page, "per_page": per_page, "objects": slice_}


@app.get("/api/museum-objects/types")
def museum_object_types():
    objs = _load_museum_objects()
    types = sorted({o.get("type", "") for o in objs if o.get("type")})
    return {"types": types}


@app.get("/api/museum-objects/{object_id}")
def museum_object_detail(object_id: str):
    objs = _load_museum_objects()
    for o in objs:
        if o.get("id") == object_id:
            cached_path = _MUSEUM_IMAGE_CACHE / f"{object_id}.jpg"
            if cached_path.exists():
                o = dict(o)
                o["image_local"] = f"/museum-resource/images/{object_id}.jpg"
            return o
    raise HTTPException(status_code=404, detail="Object not found")


# ── Movies Library ───────────────────────────────────────────────────────────
_MOVIES_PATH = Path(__file__).parent.parent / "data" / "movies.json"

@lru_cache(maxsize=1)
def _load_movies() -> list:
    if not _MOVIES_PATH.exists():
        return []
    with open(_MOVIES_PATH, encoding="utf-8") as f:
        d = json.load(f)
    return d.get("movies", d) if isinstance(d, dict) else d

@app.get("/api/movies")
def list_movies(q: str = "", genre: str = "", age_group: str = "", country: str = "", page: int = 1, per_page: int = 24):
    movies = _load_movies()
    q_lower = q.lower()
    results = []
    for m in movies:
        if q_lower and q_lower not in m.get("title", "").lower() and q_lower not in m.get("director", "").lower() and q_lower not in m.get("description", "").lower():
            continue
        if genre and genre not in m.get("genre", []):
            continue
        if age_group and m.get("age_group") != age_group:
            continue
        if country and country.lower() not in m.get("country", "").lower():
            continue
        results.append(m)
    total = len(results)
    start = (page - 1) * per_page
    return {"movies": results[start:start + per_page], "total": total}

@app.get("/api/movies/genres")
def movies_genres():
    movies = _load_movies()
    genres: set = set()
    for m in movies:
        for g in m.get("genre", []):
            genres.add(g)
    return {"genres": sorted(genres)}

@app.get("/api/movies/countries")
def movies_countries():
    movies = _load_movies()
    countries: set = set()
    for m in movies:
        for part in m.get("country", "").split("/"):
            c = part.strip()
            if c:
                countries.add(c)
    return {"countries": sorted(countries)}

@app.get("/api/movies/age-groups")
def movies_age_groups():
    movies = _load_movies()
    ages: set = set()
    for m in movies:
        if m.get("age_group"):
            ages.add(m["age_group"])
    def sort_key(a):
        try:
            return int(a.replace("+", ""))
        except Exception:
            return 99
    return {"age_groups": sorted(ages, key=sort_key)}


# ── World Literature Library ─────────────────────────────────────────────────
_WLIT_PATH = Path(__file__).parent.parent / "data" / "world_literature" / "library.json"

def _load_wlit() -> dict:
    with open(_WLIT_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/world-literature")
def world_literature_overview():
    data = _load_wlit()
    sections = []
    for key, section in data["sections"].items():
        sections.append({
            "id": key,
            "label": section["label"],
            "emoji": section["emoji"],
            "age_range": section.get("age_range", ""),
            "book_count": len(section.get("books", [])),
        })
    return {"title": data["title"], "description": data["description"], "sections": sections}

@app.get("/api/world-literature/search")
def world_literature_search(q: str = "", limit: int = 60):
    data = _load_wlit()
    q_lower = q.lower().strip()
    results = []
    if len(q_lower) >= 2:
        for section_key, section in data["sections"].items():
            for book in section.get("books", []):
                haystack = f"{book.get('title', '')} {book.get('author', '')} {book.get('genre', '')}".lower()
                if q_lower in haystack:
                    results.append({**book, "section": section_key, "section_label": section["label"]})
    return {"results": results[:limit], "total_matches": len(results)}


@app.get("/api/world-literature/{section}")
def world_literature_section(section: str):
    data = _load_wlit()
    if section not in data["sections"]:
        raise HTTPException(status_code=404, detail="Section not found")
    return {"id": section, **data["sections"][section]}

@app.get("/api/world-literature/{section}/{book_id}")
def world_literature_book(section: str, book_id: str):
    data = _load_wlit()
    if section not in data["sections"]:
        raise HTTPException(status_code=404, detail="Section not found")
    for book in data["sections"][section].get("books", []):
        if book["id"] == book_id:
            return book
    raise HTTPException(status_code=404, detail="Book not found")


# ── Biography Library ──────────────────────────────────────────────────────
_BIOGRAPHIES_PATH = Path(__file__).parent.parent / "data" / "biographies" / "biographies.json"

def _load_biographies() -> dict:
    with open(_BIOGRAPHIES_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/biographies")
def biographies_overview():
    data = _load_biographies()
    sections = []
    for key, section in data["sections"].items():
        sections.append({
            "id": key,
            "label": section["label"],
            "emoji": section["emoji"],
            "description": section.get("description", ""),
            "person_count": len(section.get("people", [])),
        })
    total = sum(s["person_count"] for s in sections)
    return {"title": data["title"], "description": data["description"], "sections": sections, "total_people": total}

@app.get("/api/biographies/search")
def biographies_search(q: str = "", limit: int = 60):
    data = _load_biographies()
    q_lower = q.lower().strip()
    results = []
    if len(q_lower) >= 2:
        for section_key, section in data["sections"].items():
            for person in section.get("people", []):
                haystack = f"{person.get('name', '')} {person.get('field', '')} {person.get('nationality', '')}".lower()
                if q_lower in haystack:
                    results.append({**person, "section": section_key, "section_label": section["label"]})
    return {"results": results[:limit], "total_matches": len(results)}


@app.get("/api/biographies/{section}")
def biographies_section(section: str):
    data = _load_biographies()
    if section not in data["sections"]:
        raise HTTPException(status_code=404, detail="Section not found")
    return {"id": section, **data["sections"][section]}

@app.get("/api/biographies/{section}/{person_id}")
def biographies_person(section: str, person_id: str):
    data = _load_biographies()
    if section not in data["sections"]:
        raise HTTPException(status_code=404, detail="Section not found")
    for person in data["sections"][section].get("people", []):
        if person["id"] == person_id:
            return person
    raise HTTPException(status_code=404, detail="Person not found")


# ── Critical Thinking Academy ─────────────────────────────────────────────────
_CT_PATH = Path(__file__).parent.parent / "data" / "critical_thinking" / "critical_thinking.json"

def _load_ct() -> dict:
    with open(_CT_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/critical-thinking")
def critical_thinking_overview():
    data = _load_ct()
    modules = []
    for key, mod in data["modules"].items():
        modules.append({
            "id": key,
            "label": mod["label"],
            "emoji": mod["emoji"],
            "description": mod["description"],
            "lesson_count": len(mod.get("lessons", [])),
        })
    return {"title": data["title"], "description": data["description"], "modules": modules}

@app.get("/api/critical-thinking/{module_id}")
def critical_thinking_module(module_id: str):
    data = _load_ct()
    if module_id not in data["modules"]:
        raise HTTPException(status_code=404, detail="Module not found")
    return {"id": module_id, **data["modules"][module_id]}

@app.get("/api/critical-thinking/{module_id}/{lesson_id}")
def critical_thinking_lesson(module_id: str, lesson_id: str):
    data = _load_ct()
    if module_id not in data["modules"]:
        raise HTTPException(status_code=404, detail="Module not found")
    for lesson in data["modules"][module_id].get("lessons", []):
        if lesson["id"] == lesson_id:
            return lesson
    raise HTTPException(status_code=404, detail="Lesson not found")

# ── Survival Skills ──────────────────────────────────────────────────────────
_SURVIVAL_PATH = Path(__file__).parent.parent / "data" / "survival_skills" / "survival_skills.json"

def _load_survival() -> dict:
    with open(_SURVIVAL_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/survival-skills")
def survival_overview():
    data = _load_survival()
    cats = []
    for cid, cat in data["categories"].items():
        if isinstance(cat, dict):
            label = cat.get("label", cid.replace("_", " ").title())
            emoji = cat.get("emoji", "🛡️")
            skills_list = cat.get("skills", [])
        else:
            label = cid.replace("_", " ").title()
            emoji = "🛡️"
            skills_list = cat
        cats.append({"id": cid, "label": label, "emoji": emoji, "skill_count": len(skills_list)})
    return {"title": data["title"], "description": data["description"], "categories": cats}

@app.get("/api/survival-skills/{category}")
def survival_category(category: str):
    data = _load_survival()
    cat = data["categories"].get(category)
    if cat is None:
        raise HTTPException(status_code=404, detail="Category not found")
    if isinstance(cat, list):
        label = category.replace("_", " ").title()
        return {"id": category, "label": label, "skills": cat}
    return cat

@app.get("/api/survival-skills/{category}/{skill_name}")
def survival_skill(category: str, skill_name: str):
    data = _load_survival()
    cat = data["categories"].get(category)
    if cat is None:
        raise HTTPException(status_code=404, detail="Category not found")
    skills_list = cat if isinstance(cat, list) else cat.get("skills", [])
    for s in skills_list:
        if s.get("id") == skill_name or s.get("name", "").lower().replace(" ", "_") == skill_name:
            return s
    raise HTTPException(status_code=404, detail="Skill not found")


# ── Environmental Science ────────────────────────────────────────────────────
_ENV_PATH = Path(__file__).parent.parent / "data" / "environmental_science" / "environmental_science.json"

def _load_env() -> dict:
    with open(_ENV_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/environmental-science")
def env_overview():
    data = _load_env()
    units = []
    for uid, unit in data["units"].items():
        units.append({"id": uid, "label": unit["label"], "emoji": unit["emoji"],
                      "topic_count": len(unit["topics"])})
    return {"title": data["title"], "description": data["description"], "units": units}

@app.get("/api/environmental-science/{unit}")
def env_unit(unit: str):
    data = _load_env()
    u = data["units"].get(unit)
    if not u:
        raise HTTPException(status_code=404, detail="Unit not found")
    return u

@app.get("/api/environmental-science/{unit}/{topic_id}")
def env_topic(unit: str, topic_id: str):
    data = _load_env()
    u = data["units"].get(unit)
    if not u:
        raise HTTPException(status_code=404, detail="Unit not found")
    for topic in u["topics"]:
        if topic["id"] == topic_id:
            return topic
    raise HTTPException(status_code=404, detail="Topic not found")


# ── World Politics ────────────────────────────────────────────────────────────
_WPOL_PATH = Path(__file__).parent.parent / "data" / "world_politics" / "world_politics.json"

def _load_wpol() -> dict:
    with open(_WPOL_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/world-politics")
def world_politics_overview():
    data = _load_wpol()
    modules = []
    for mid, mod in data["modules"].items():
        lesson_count = len(mod.get("lessons", [])) or len(mod.get("countries", {}))
        modules.append({"id": mid, "label": mod["label"], "emoji": mod["emoji"],
                        "description": mod["description"],
                        "lesson_count": lesson_count})
    return {"title": data["title"], "description": data["description"],
            "disclaimer": data.get("disclaimer", ""), "modules": modules}

@app.get("/api/world-politics/countries/{country_id}")
def world_politics_country(country_id: str):
    data = _load_wpol()
    countries = data["modules"].get("country_profiles", {}).get("countries", {})
    country = countries.get(country_id)
    if not country:
        raise HTTPException(status_code=404, detail=f"Country '{country_id}' not found")
    return country

@app.get("/api/world-politics/countries")
def world_politics_countries():
    data = _load_wpol()
    countries = data["modules"].get("country_profiles", {}).get("countries", {})
    return {"countries": [{"id": k, "name": v["name"], "flag": v.get("flag", "🌍"),
                           "government_type": v["government_type"]} for k, v in countries.items()]}

@app.get("/api/world-politics/{module_id}")
def world_politics_module(module_id: str):
    data = _load_wpol()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    return mod

@app.get("/api/world-politics/{module_id}/{lesson_id}")
def world_politics_lesson(module_id: str, lesson_id: str):
    data = _load_wpol()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    for lesson in mod.get("lessons", []):
        if lesson["id"] == lesson_id:
            return lesson
    raise HTTPException(status_code=404, detail="Lesson not found")


# ── World Religions ───────────────────────────────────────────────────────────
_RELIGIONS_PATH = Path(__file__).parent.parent / "data" / "world_religions" / "world_religions.json"

def _load_religions() -> dict:
    with open(_RELIGIONS_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/world-religions")
def world_religions_overview():
    data = _load_religions()
    religions = []
    for rid, rel in data["religions"].items():
        religions.append({"id": rid, "name": rel["name"], "emoji": rel.get("emoji", "🕌"),
                          "adherents_approx": rel.get("adherents_approx", ""),
                          "founded": rel.get("founded", ""),
                          "origin": rel.get("origin", ""),
                          "summary": rel.get("summary", ""),
                          "lesson_count": len(rel.get("lessons", []))})
    return {"title": data["title"], "description": data["description"],
            "disclaimer": data.get("disclaimer", ""), "religions": religions}

@app.get("/api/world-religions/{religion_id}")
def world_religion_detail(religion_id: str):
    data = _load_religions()
    rel = data["religions"].get(religion_id)
    if not rel:
        raise HTTPException(status_code=404, detail=f"Religion '{religion_id}' not found")
    return rel


# ── Health Education ──────────────────────────────────────────────────────────
_HEALTH_PATH = Path(__file__).parent.parent / "data" / "health_education" / "health_education.json"

def _load_health() -> dict:
    with open(_HEALTH_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/health-education")
def health_overview():
    data = _load_health()
    units = []
    for uid, unit in data["units"].items():
        units.append({"id": uid, "label": unit["label"], "emoji": unit["emoji"],
                      "topic_count": len(unit["topics"])})
    return {"title": data["title"], "description": data["description"], "units": units}

@app.get("/api/health-education/{unit}")
def health_unit(unit: str):
    data = _load_health()
    u = data["units"].get(unit)
    if not u:
        raise HTTPException(status_code=404, detail="Unit not found")
    return u

@app.get("/api/health-education/{unit}/{topic_id}")
def health_topic(unit: str, topic_id: str):
    data = _load_health()
    u = data["units"].get(unit)
    if not u:
        raise HTTPException(status_code=404, detail="Unit not found")
    for topic in u["topics"]:
        if topic["id"] == topic_id:
            return topic
    raise HTTPException(status_code=404, detail="Topic not found")


# ── Business Studies ──────────────────────────────────────────────────────────
_BIZ_PATH = Path(__file__).parent.parent / "data" / "business_studies" / "business_studies.json"

def _load_biz() -> dict:
    with open(_BIZ_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/business-studies")
def business_overview():
    data = _load_biz()
    modules = []
    for mid, mod in data["modules"].items():
        modules.append({"id": mid, "label": mod["label"], "emoji": mod["emoji"],
                        "description": mod["description"],
                        "lesson_count": len(mod["lessons"])})
    return {"title": data["title"], "description": data["description"], "modules": modules}

@app.get("/api/business-studies/{module_id}")
def business_module(module_id: str):
    data = _load_biz()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    return mod

@app.get("/api/business-studies/{module_id}/{lesson_id}")
def business_lesson(module_id: str, lesson_id: str):
    data = _load_biz()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    for lesson in mod["lessons"]:
        if lesson["id"] == lesson_id:
            return lesson
    raise HTTPException(status_code=404, detail="Lesson not found")


# ── Attendance Tracking ───────────────────────────────────────────────────────
def _load_att(child: str) -> list:
    return get_attendance_records(child)

def _save_att(child: str, records: list):
    return save_attendance_records(child, records)

@app.get("/api/parent/attendance/{child}")
def get_attendance(child: str):
    _require_child(child)
    return {"child": child, "records": _load_att(child)}

@app.post("/api/parent/attendance/{child}")
def add_attendance(child: str, body: dict):
    _require_child(child)
    records = _load_att(child)
    from datetime import date as _date
    record = {
        "date": body.get("date", str(_date.today())),
        "status": body.get("status", "present"),  # present | absent | late | excused
        "note": body.get("note", ""),
    }
    records.append(record)
    _save_att(child, records)
    return {"ok": True, "record": record}

@app.get("/api/parent/attendance/{child}/summary")
def attendance_summary(child: str):
    _require_child(child)
    records = _load_att(child)
    counts = {"present": 0, "absent": 0, "late": 0, "excused": 0}
    for r in records:
        counts[r.get("status", "present")] = counts.get(r.get("status", "present"), 0) + 1
    total = len(records)
    rate = round(counts["present"] / total * 100, 1) if total else 0
    return {"child": child, "total_days": total, "attendance_rate": rate, "counts": counts}

@app.delete("/api/parent/attendance/{child}/{date}")
def delete_attendance(child: str, date: str):
    _require_child(child)
    records = [r for r in _load_att(child) if r["date"] != date]
    _save_att(child, records)
    return {"ok": True}


# ── Civics ────────────────────────────────────────────────────────────────────
_CIVICS_PATH = Path(__file__).parent.parent / "data" / "civics" / "civics.json"

def _load_civics() -> dict:
    with open(_CIVICS_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/civics")
def civics_overview():
    data = _load_civics()
    modules = []
    for mid, mod in data["modules"].items():
        modules.append({"id": mid, "label": mod["label"], "emoji": mod["emoji"],
                        "description": mod["description"],
                        "lesson_count": len(mod["lessons"])})
    return {"title": data["title"], "description": data["description"], "modules": modules}

@app.get("/api/civics/{module_id}")
def civics_module(module_id: str):
    data = _load_civics()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    return mod

@app.get("/api/civics/{module_id}/{lesson_id}")
def civics_lesson(module_id: str, lesson_id: str):
    data = _load_civics()
    mod = data["modules"].get(module_id)
    if not mod:
        raise HTTPException(status_code=404, detail="Module not found")
    for lesson in mod["lessons"]:
        if lesson["id"] == lesson_id:
            return lesson
    raise HTTPException(status_code=404, detail="Lesson not found")

# ── Song Centre ───────────────────────────────────────────────────────────────
_SONGS_PATH = Path(__file__).parent.parent / "data" / "song_centre" / "songs.json"

def _load_songs():
    with open(_SONGS_PATH, encoding="utf-8") as f:
        return json.load(f)

@app.get("/api/songs")
def songs_overview():
    data = _load_songs()
    cards = []
    for song in data["songs"]:
        card = {
            key: song[key]
            for key in (
                "id", "title", "artist", "year", "genre", "origin_country",
                "language", "decade", "suitable_for_ages", "tags", "links",
            )
        }
        for optional in ("chart_rank", "verified_views", "view_count_checked_at"):
            if optional in song:
                card[optional] = song[optional]
        cards.append(card)
    return {
        "title": data["title"],
        "description": data["description"],
        "total": data["total"],
        "genres": data["genres"],
        "decades": data["decades"],
        "songs": cards,
    }

@app.get("/api/songs/genre/{genre}")
def songs_by_genre(genre: str):
    data = _load_songs()
    genre_lower = genre.lower()
    matches = [s for s in data["songs"]
               if any(genre_lower in g.lower() for g in s.get("genre", []))]
    return {"genre": genre, "songs": matches, "count": len(matches)}

@app.get("/api/songs/decade/{decade}")
def songs_by_decade(decade: str):
    data = _load_songs()
    matches = [s for s in data["songs"] if s.get("decade","") == decade]
    return {"decade": decade, "songs": matches, "count": len(matches)}

@app.get("/api/songs/{song_id}")
def song_detail(song_id: str):
    data = _load_songs()
    for s in data["songs"]:
        if s["id"] == song_id:
            return s
    raise HTTPException(status_code=404, detail="Song not found")


def _yt(q: str) -> str:
    return "https://www.youtube.com/results?search_query=" + q.replace(" ", "+")

def _wiki(q: str) -> str:
    return "https://en.wikipedia.org/w/index.php?search=" + q.replace(" ", "+")

def _lyr(q: str) -> str:
    return "https://www.google.com/search?q=" + q.replace(" ", "+") + "+lyrics"

def _url_valid(url) -> bool:
    """Return True if url is a non-empty http(s) string with no obvious placeholder."""
    if not isinstance(url, str) or not url.strip():
        return False
    u = url.strip()
    if not (u.startswith("http://") or u.startswith("https://")):
        return False
    BAD = ("example.com", "placeholder", "PLACEHOLDER", "TODO", "javascript:",
           "watch?v=XXXX", "watch?v=undefined", "watch?v=VIDEO_ID")
    return not any(b in u for b in BAD)

def _song_links_ok(links: dict) -> bool:
    yt   = links.get("youtube_search", "")
    wiki = links.get("wiki_search", "")
    lyr  = links.get("lyrics_search", "")
    return (
        _url_valid(yt)   and "results?search_query=" in yt and "watch?v=" not in yt and
        _url_valid(wiki) and ("search=" in wiki or "wikipedia.org/wiki/" in wiki) and
        _url_valid(lyr)  and "google.com/search" in lyr
    )

def _fix_songs(songs_path: Path) -> dict:
    """Fix song search links and return stats."""
    data = json.loads(songs_path.read_text(encoding="utf-8"))
    fixed = []
    for song in data["songs"]:
        links = song.get("links") or {}
        q = f"{song['title']} {song['artist']}"
        changed = False
        if not (_url_valid(links.get("youtube_search","")) and "results?search_query=" in links.get("youtube_search","")):
            links["youtube_search"] = _yt(q + " official"); changed = True
        if not (_url_valid(links.get("wiki_search","")) and "search=" in links.get("wiki_search","")):
            links["wiki_search"] = _wiki(q); changed = True
        if not (_url_valid(links.get("lyrics_search","")) and "google.com/search" in links.get("lyrics_search","")):
            links["lyrics_search"] = _lyr(q); changed = True
        if changed:
            song["links"] = links
            fixed.append(song.get("title", "?"))
    if fixed:
        data["total"] = len(data["songs"])
        songs_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"checked": len(data["songs"]), "fixed": len(fixed), "fixed_items": fixed[:10]}


def _scan_json_urls(path: Path) -> dict:
    """
    Walk a JSON file, collect all string values that look like they should be URLs
    (keys containing 'url' or 'link'). Report invalid ones; we cannot auto-fix
    these since we don't know the correct replacement URL for curated content.
    """
    URL_KEYS = {"url", "link", "video_link", "resource_link", "text_link", "audiourl", "videourl", "channel_url"}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {"checked": 0, "broken": 0, "broken_items": []}

    checked = 0
    broken = []

    def walk(obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                if k.lower().replace("_", "") in URL_KEYS or k.lower().endswith("url") or k.lower().endswith("link"):
                    if isinstance(v, str) and v:
                        checked_here = True
                        if not _url_valid(v):
                            title = obj.get("title") or obj.get("name") or obj.get("id") or "?"
                            broken.append({"field": k, "value": v[:80], "context": str(title)[:60]})
                walk(v)
        elif isinstance(obj, list):
            for item in obj:
                walk(item)

    # patch nonlocal counts
    def walk_counted(obj):
        nonlocal checked
        if isinstance(obj, dict):
            for k, v in obj.items():
                kl = k.lower().replace("_", "")
                if kl in URL_KEYS or k.lower().endswith("url") or k.lower().endswith("link"):
                    if isinstance(v, str) and v:
                        checked += 1
                        if not _url_valid(v):
                            title = obj.get("title") or obj.get("name") or obj.get("id") or "?"
                            broken.append({"field": k, "value": v[:80], "context": str(title)[:60]})
                walk_counted(v)
        elif isinstance(obj, list):
            for item in obj:
                walk_counted(item)

    walk_counted(data)
    return {"checked": checked, "broken": len(broken), "broken_items": broken[:5]}


# Keep old route for backwards compat (delegates to global)
@app.post("/api/songs/refresh-links")
def refresh_song_links():
    result = _fix_songs(_SONGS_PATH)
    return {
        "total_checked": result["checked"],
        "fixed": result["fixed"],
        "ok": result["checked"] - result["fixed"],
        "fixed_songs": [{"title": t} for t in result["fixed_items"]],
    }


_DATA_ROOT    = Path(__file__).parent.parent / "data"
_SYLLABUS_ROOT = Path(__file__).parent.parent / "syllabus"
_LINK_AUDIT_PATH = _DATA_ROOT / "link_audit_status.json"


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _load_link_audit() -> dict:
    if not _LINK_AUDIT_PATH.exists():
        return {"urls": {}, "runs": []}
    try:
        return json.loads(_LINK_AUDIT_PATH.read_text(encoding="utf-8"))
    except Exception:
        return {"urls": {}, "runs": []}


def _save_link_audit(data: dict) -> None:
    _LINK_AUDIT_PATH.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def _verify_url_server_side(url: str, timeout: float = 4.0) -> dict:
    if not _url_valid(url):
        return {"review_status": "broken", "http_status": None, "error": "invalid-url"}
    headers = {"User-Agent": "EduAI-LinkAudit/1.0"}
    for method in ("HEAD", "GET"):
        try:
            req = Request(url, method=method, headers=headers)
            with urlopen(req, timeout=timeout) as resp:
                status = getattr(resp, "status", 200)
                return {
                    "review_status": "ok" if status < 400 else "broken",
                    "http_status": status,
                    "error": None,
                }
        except HTTPError as exc:
            if method == "HEAD" and exc.code in {403, 405}:
                continue
            return {"review_status": "broken", "http_status": exc.code, "error": str(exc.reason)}
        except URLError as exc:
            if method == "HEAD":
                continue
            return {"review_status": "broken", "http_status": None, "error": str(exc.reason)}
        except Exception as exc:
            if method == "HEAD":
                continue
            return {"review_status": "broken", "http_status": None, "error": str(exc)}
    return {"review_status": "broken", "http_status": None, "error": "unverified"}


@app.get("/api/link-audit")
def get_link_audit():
    data = _load_link_audit()
    urls = data.get("urls", {})
    counts = {"ok": 0, "broken": 0, "unverified": 0}
    for record in urls.values():
        status = record.get("review_status", "unverified")
        counts[status if status in counts else "unverified"] += 1
    return {
        "total": len(urls),
        "counts": counts,
        "last_run": (data.get("runs") or [None])[-1],
        "urls": list(urls.values()),
    }


@app.post("/api/link-audit/run")
def run_link_audit(limit_per_file: int = 50, max_urls: int = 200, verify: bool = False):
    records = _collect_all_urls(limit_per_file)[:max(1, min(max_urls, 1000))]
    audit = _load_link_audit()
    urls = audit.setdefault("urls", {})
    checked_at = _now_iso()
    counts = {"ok": 0, "broken": 0, "unverified": 0}

    for record in records:
        url = record["url"]
        status = _verify_url_server_side(url) if verify else {
            "review_status": "unverified" if _url_valid(url) else "broken",
            "http_status": None,
            "error": None if _url_valid(url) else "invalid-url",
        }
        saved = {
            **record,
            **status,
            "checked_at": checked_at,
            "verification": "server" if verify else "structural",
        }
        urls[url] = saved
        review_status = saved.get("review_status", "unverified")
        counts[review_status if review_status in counts else "unverified"] += 1

    run = {
        "checked_at": checked_at,
        "checked": len(records),
        "verify": verify,
        "counts": counts,
    }
    runs = audit.setdefault("runs", [])
    runs.append(run)
    audit["runs"] = runs[-20:]
    _save_link_audit(audit)
    return {"stored": len(urls), **run}


@app.post("/api/content-store/reindex")
def reindex_content_store():
    return content_store.reindex()


@app.get("/api/content-store/search")
def search_content_store(q: str = "", limit: int = 50):
    if not q.strip():
        return {"total": 0, "items": []}
    return content_store.search(q.strip(), limit)


@app.post("/api/refresh-all-links")
def refresh_all_links():
    """
    Global link health-check + auto-repair for all content files.
    Songs: auto-fix search links.
    All other JSON: report broken URLs (cannot auto-fix curated content URLs).
    """
    report = {}
    total_checked = 0
    total_fixed = 0
    total_broken = 0

    # ── Songs ──────────────────────────────────────────────────────────────────
    songs_result = _fix_songs(_SONGS_PATH)
    report["Song Centre"] = {
        "checked": songs_result["checked"],
        "fixed":   songs_result["fixed"],
        "broken":  0,
        "note": f"Auto-fixed {songs_result['fixed']} search link(s)." if songs_result["fixed"] else "All search links healthy.",
    }
    total_checked += songs_result["checked"]
    total_fixed   += songs_result["fixed"]

    # ── Syllabus (grades 1-10) ─────────────────────────────────────────────────
    for gf in sorted(_SYLLABUS_ROOT.glob("grade*.json")):
        r = _scan_json_urls(gf)
        label = gf.stem.replace("grade", "Grade ")
        report[label] = {"checked": r["checked"], "fixed": 0, "broken": r["broken"],
                         "broken_items": r["broken_items"]}
        total_checked += r["checked"]
        total_broken  += r["broken"]

    # ── All other data JSONs ───────────────────────────────────────────────────
    SKIP = {"progress_", "activity_", "attendance_", "users.json", "songs.json"}
    for jf in sorted(_DATA_ROOT.rglob("*.json")):
        if any(s in jf.name for s in SKIP):
            continue
        r = _scan_json_urls(jf)
        if r["checked"] == 0:
            continue
        # Use parent folder + filename as label
        parts = jf.relative_to(_DATA_ROOT).parts
        label = " / ".join(parts).replace(".json", "").replace("_", " ").title()
        report[label] = {"checked": r["checked"], "fixed": 0, "broken": r["broken"],
                         "broken_items": r["broken_items"]}
        total_checked += r["checked"]
        total_broken  += r["broken"]

    return {
        "total_checked": total_checked,
        "total_fixed":   total_fixed,
        "total_broken":  total_broken,
        "categories":    len(report),
        "report":        report,
    }


# ─── Browser-side live link-check support ────────────────────────────────────

def _collect_all_urls(limit_per_file: int = 300) -> list[dict]:
    """
    Walk every content JSON and collect URL records for browser-side checking.
    Each record: {url, title, source, file, json_key_path, category}
    We cap per file so the list stays manageable for a browser scan.
    """
    URL_KEYS = {"url", "link", "video_link", "resource_link", "text_link",
                "audiourl", "videourl", "channel_url",
                "youtube_search", "wiki_search", "lyrics_search"}
    SKIP_FILES = {"progress_", "activity_", "attendance_", "users.json"}
    records: list[dict] = []

    def walk(obj, path: list, file_label: str, category: str, count: list):
        if count[0] >= limit_per_file:
            return
        if isinstance(obj, dict):
            title  = obj.get("title") or obj.get("name") or obj.get("id") or ""
            source = obj.get("source") or obj.get("artist") or ""
            for k, v in obj.items():
                if k.lower().replace("_","") in URL_KEYS or k.lower().endswith("url") or k.lower().endswith("link"):
                    if isinstance(v, str) and v.startswith("http"):
                        records.append({
                            "url":      v,
                            "title":    str(title)[:120],
                            "source":   str(source)[:80],
                            "file":     file_label,
                            "key":      k,
                            "category": category,
                        })
                        count[0] += 1
                        if count[0] >= limit_per_file:
                            return
            for k, v in obj.items():
                walk(v, path + [k], file_label, category, count)
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                walk(item, path + [str(i)], file_label, category, count)

    def process(path: Path, category: str):
        if any(s in path.name for s in SKIP_FILES):
            return
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return
        label = str(path.relative_to(path.parent.parent))
        walk(data, [], label, category, [0])

    # Songs
    process(_SONGS_PATH, "Song Centre")
    # Syllabus
    for gf in sorted(_SYLLABUS_ROOT.glob("grade*.json")):
        process(gf, gf.stem.replace("grade", "Grade "))
    # Other data
    for jf in sorted(_DATA_ROOT.rglob("*.json")):
        if any(s in jf.name for s in SKIP_FILES) or "songs.json" in jf.name:
            continue
        parts = jf.relative_to(_DATA_ROOT).parts
        cat = " / ".join(parts).replace(".json","").replace("_"," ").title()
        process(jf, cat)

    return records


@app.get("/api/get-all-links")
def get_all_links(limit_per_file: int = 200):
    """Return all content URLs for browser-side live-checking."""
    records = _collect_all_urls(limit_per_file)
    return {"total": len(records), "urls": records}


class LinkFix(BaseModel):
    url: str         # original broken URL
    replacement: str # new URL to substitute in
    file: str        # relative file path (e.g. syllabus/grade1.json)
    key: str         # JSON field name (e.g. "url")
    title: str       # resource title for logging

class LinkFixBatch(BaseModel):
    fixes: list[LinkFix]

@app.post("/api/apply-link-fixes")
def apply_link_fixes(body: LinkFixBatch):
    """
    Receive a list of {url, replacement, file, key} from the browser after
    live-checking, and apply the substitutions to the JSON files.
    """
    if not body.fixes:
        return {"applied": 0}

    # Group fixes by file
    by_file: dict[str, list[LinkFix]] = {}
    for fix in body.fixes:
        by_file.setdefault(fix.file, []).append(fix)

    applied = 0
    errors: list[str] = []

    for rel_path, fixes in by_file.items():
        # Resolve path relative to backend root
        backend_root = Path(__file__).parent.parent
        full_path = (backend_root / rel_path).resolve()
        # Safety: must stay within backend directory
        try:
            full_path.relative_to(backend_root.resolve())
        except ValueError:
            errors.append(f"Unsafe path rejected: {rel_path}")
            continue
        if not full_path.exists():
            errors.append(f"File not found: {rel_path}")
            continue

        try:
            text = full_path.read_text(encoding="utf-8")
            for fix in fixes:
                if not fix.replacement.startswith("http"):
                    continue
                # Simple string substitution — safe because URLs are unique values
                new_text = text.replace(
                    json.dumps(fix.url),        # "https://old..."
                    json.dumps(fix.replacement) # "https://new..."
                )
                if new_text != text:
                    text = new_text
                    applied += 1
            # Validate JSON is still valid before writing
            json.loads(text)
            full_path.write_text(text, encoding="utf-8")
        except Exception as e:
            errors.append(f"{rel_path}: {e}")

    return {"applied": applied, "errors": errors}


# ── Code execution endpoint ───────────────────────────────────────────────────

class CodeRunRequest(BaseModel):
    # "cpp", "fortran", "sql", "java", "c", "go", "rust", "php", "ruby", "csharp", "perl", "r"
    language: str
    code: str

_TIMEOUT = 10  # seconds

def _run_subprocess(cmd: list[str], input_text: str | None = None, env: dict | None = None) -> str:
    try:
        result = subprocess.run(
            cmd,
            input=input_text,
            capture_output=True,
            text=True,
            timeout=_TIMEOUT,
            env=env,
        )
        out = result.stdout or ""
        err = result.stderr or ""
        return (out + err).strip() or "(no output)"
    except subprocess.TimeoutExpired:
        return f"Error: execution timed out after {_TIMEOUT} seconds"
    except Exception as e:
        return f"Error: {e}"


@app.post("/api/run-code")
async def run_code(req: CodeRunRequest):
    lang = req.language.lower()
    code = req.code

    if lang == "sql":
        # Run SQL against an in-memory SQLite database
        try:
            conn = _sqlite3.connect(":memory:")
            conn.row_factory = _sqlite3.Row
            cur = conn.cursor()
            lines = []
            # Execute each statement; collect SELECT results
            for stmt in code.split(";"):
                stmt = stmt.strip()
                if not stmt:
                    continue
                try:
                    cur.execute(stmt)
                    if stmt.upper().startswith("SELECT"):
                        rows = cur.fetchall()
                        if rows:
                            cols = [d[0] for d in cur.description]
                            lines.append(" | ".join(cols))
                            lines.append("-" * (sum(len(c) for c in cols) + 3 * (len(cols) - 1)))
                            for row in rows:
                                lines.append(" | ".join(str(v) for v in row))
                        else:
                            lines.append("(no rows)")
                    else:
                        lines.append(f"OK ({cur.rowcount} row(s) affected)")
                except _sqlite3.Error as e:
                    lines.append(f"SQL Error: {e}")
            conn.close()
            output = "\n".join(lines) if lines else "(no output)"
        except Exception as e:
            output = f"Error: {e}"
        return {"output": output}

    with tempfile.TemporaryDirectory() as tmpdir:
        if lang == "cpp":
            src = os.path.join(tmpdir, "main.cpp")
            exe = os.path.join(tmpdir, "main")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["g++", "-o", exe, src, "-std=c++17"])
            if not os.path.exists(exe):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess([exe])}

        if lang == "fortran":
            src = os.path.join(tmpdir, "main.f90")
            exe = os.path.join(tmpdir, "main")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["gfortran", "-o", exe, src])
            if not os.path.exists(exe):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess([exe])}

        if lang == "c":
            src = os.path.join(tmpdir, "main.c")
            exe = os.path.join(tmpdir, "main")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["gcc", "-o", exe, src])
            if not os.path.exists(exe):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess([exe])}

        if lang == "java":
            # Expects the code to declare `public class Main`.
            src = os.path.join(tmpdir, "Main.java")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["javac", src])
            if not os.path.exists(os.path.join(tmpdir, "Main.class")):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess(["java", "-cp", tmpdir, "Main"])}

        if lang == "csharp":
            src = os.path.join(tmpdir, "main.cs")
            exe = os.path.join(tmpdir, "main.exe")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["mcs", f"-out:{exe}", src])
            if not os.path.exists(exe):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess(["mono", exe])}

        if lang == "go":
            src = os.path.join(tmpdir, "main.go")
            Path(src).write_text(code)
            # Give Go its own writable cache/home inside the tempdir so the
            # sandboxed process doesn't depend on a shared $HOME build cache.
            go_env = {
                **os.environ,
                "HOME": tmpdir,
                "GOCACHE": os.path.join(tmpdir, "gocache"),
                "GOPATH": os.path.join(tmpdir, "gopath"),
            }
            return {"output": _run_subprocess(["go", "run", src], env=go_env)}

        if lang == "rust":
            src = os.path.join(tmpdir, "main.rs")
            exe = os.path.join(tmpdir, "main")
            Path(src).write_text(code)
            compile_out = _run_subprocess(["rustc", "-o", exe, src])
            if not os.path.exists(exe):
                return {"output": f"Compile error:\n{compile_out}"}
            return {"output": _run_subprocess([exe])}

        if lang == "php":
            src = os.path.join(tmpdir, "main.php")
            Path(src).write_text(code)
            return {"output": _run_subprocess(["php", src])}

        if lang == "ruby":
            src = os.path.join(tmpdir, "main.rb")
            Path(src).write_text(code)
            return {"output": _run_subprocess(["ruby", src])}

        if lang == "perl":
            src = os.path.join(tmpdir, "main.pl")
            Path(src).write_text(code)
            return {"output": _run_subprocess(["perl", src])}

        if lang == "r":
            src = os.path.join(tmpdir, "main.R")
            Path(src).write_text(code)
            return {"output": _run_subprocess(["Rscript", src])}

    return {"output": f"Unsupported language: {lang}"}


# ── Serve React frontend build ────────────────────────────────────────────────
# Mount the compiled React app so the backend serves the frontend at the same
# origin — this eliminates the need for a dev proxy in production deployments.
from fastapi.responses import FileResponse  # noqa: E402

_frontend_dist = Path(__file__).parent.parent.parent / "frontend" / "dist"

if _frontend_dist.exists():
    # Serve static assets (JS, CSS, images) under /assets
    app.mount("/assets", StaticFiles(directory=str(_frontend_dist / "assets")), name="assets")

    @app.get("/sw.js")
    async def service_worker():
        return FileResponse(str(_frontend_dist / "sw.js"), media_type="application/javascript")

    @app.get("/workbox-{rest:path}")
    async def workbox(rest: str):
        return FileResponse(str(_frontend_dist / f"workbox-{rest}"))

    # Catch-all: serve index.html for any non-API route (React SPA routing)
    @app.get("/{full_path:path}")
    async def spa_fallback(full_path: str):
        index = _frontend_dist / "index.html"
        if index.exists():
            return FileResponse(str(index))
        return {"error": "Frontend not built. Run: cd frontend && npm run build"}
